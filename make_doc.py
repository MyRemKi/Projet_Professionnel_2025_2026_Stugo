"""
Génère INSTALLATION.docx — procédure d'installation StuGo CO2 Explorer v7.0
Style : sobre, lisible, sans boîtes colorées — écrit comme un humain.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


doc = Document()
doc.core_properties.author = "Remi Kalkan"
doc.core_properties.last_modified_by = "Remi Kalkan"

# Marges raisonnables
for s in doc.sections:
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.5)

# Police par défaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


# ── Helpers ──────────────────────────────────────────────────────────────────

BLUE = RGBColor(0x1A, 0x56, 0x9A)   # un bleu classique, discret
GRAY = RGBColor(0x55, 0x55, 0x55)


def para(text='', align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def title_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE
    return p


def subtitle_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    return p


def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    # Filet bas simple sous le titre de section
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A569A')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def body(text, before=0, after=5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


def step(num, heading, *lines):
    # Numéro + titre sur une ligne
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"Étape {num}  —  ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(heading)
    r2.bold = True
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    # Lignes de détail
    for line in lines:
        body(line, before=0, after=3, indent=True)


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run("Note : " + text)
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p


def hrule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# CONTENU
# ─────────────────────────────────────────────────────────────────────────────

title_line("StuGo CO2 Explorer")
subtitle_line("Procédure d'installation  ·  Version 7.0  ·  Windows")

hrule()

body(
    "Ce document explique comment installer le logiciel sur votre ordinateur. "
    "Aucune connaissance technique n'est nécessaire. Suivez simplement les étapes "
    "dans l'ordre, cela prend entre 2 et 5 minutes.",
    before=12, after=6
)


# ── Ce qu'il faut avant de commencer ─────────────────────────────────────────
section_title("Ce qu'il faut avant de commencer")

body("Vérifiez que votre ordinateur correspond à ces quelques points :", after=4)
bullet("Windows 10 ou Windows 11, en version 64 bits")
bullet("Environ 500 Mo d'espace libre sur le disque")
bullet("Aucun droit administrateur n'est nécessaire — l'installation se fait dans votre dossier personnel")

body(
    "Si vous ne savez pas quelle version de Windows vous avez, appuyez sur "
    "la touche Windows + R, tapez winver et appuyez sur Entrée.",
    before=6, after=4
)


# ── Installation ──────────────────────────────────────────────────────────────
section_title("L'installation, étape par étape")

step("1", "Lancer le fichier d'installation",
    "Double-cliquez sur le fichier  StuGoCO2_Setup_v7.0.exe  pour démarrer l'installation.",
    "Une fenêtre s'ouvre automatiquement.")
note(
    "Windows peut afficher un message « Éditeur inconnu » ou une alerte de sécurité. "
    "C'est normal pour un logiciel sans certificat commercial. Cliquez sur "
    "Exécuter quand même pour continuer."
)

step("2", "Choisir la langue",
    "Une petite fenêtre vous demande de choisir la langue.",
    "Sélectionnez Français, puis cliquez sur OK.")

step("3", "Lire les informations d'installation",
    "Un écran résume ce qui va être installé et où vos données seront stockées.",
    "Prenez le temps de lire, puis cliquez sur Suivant.")

step("4", "Confirmer le dossier d'installation",
    "L'installeur propose automatiquement un dossier dans votre espace personnel :",
    "   C:\\Users\\<votre nom>\\AppData\\Local\\Programs\\StuGo CO2 Explorer\\",
    "Vous n'avez rien à modifier. Cliquez simplement sur Suivant.")
note(
    "Le logiciel s'installe dans votre espace personnel, pas dans les fichiers "
    "système de Windows. Aucun mot de passe administrateur ne sera demandé."
)

step("5", "Garder le nom du groupe Menu Démarrer",
    "Le nom « StuGo CO2 Explorer » est déjà rempli.",
    "Cliquez sur Suivant sans rien changer.")

step("6", "Choisir si vous voulez un raccourci sur le Bureau",
    "Deux options sont disponibles (désactivées par défaut) :",
    "   - Créer un raccourci sur le Bureau pour l'application principale",
    "   - Créer un raccourci sur le Bureau pour l'outil Admin Logs",
    "Cochez celles qui vous conviennent, puis cliquez sur Suivant.")

step("7", "Lancer l'installation",
    "Un récapitulatif s'affiche avec tout ce que vous avez choisi.",
    "Cliquez sur Installer.",
    "Une barre de progression s'affiche. L'opération dure entre 1 et 3 minutes.")
note("Ne fermez pas la fenêtre pendant l'installation.")

step("8", "Terminer",
    "Quand la barre est complète, l'installation est finie.",
    "L'option Lancer StuGo CO2 Explorer est cochée par défaut.",
    "Cliquez sur Terminer — le logiciel démarre.")
note(
    "Lors du tout premier démarrage, le logiciel prend 5 à 15 secondes "
    "supplémentaires. C'est normal : il prépare son environnement. "
    "Les fois suivantes, il démarre plus vite."
)


# ── Où sont les données ────────────────────────────────────────────────────────
section_title("Où sont stockées vos données ?")

body(
    "Toutes vos données sont dans votre dossier personnel. Elles ne sont "
    "jamais effacées lors d'une mise à jour du logiciel.",
    after=6
)

bullet("Sessions et filtres en cours  →  sauvegardés automatiquement")
bullet("Presets de couleurs personnalisés  →  dossier PresetColors")
bullet("Préférences (thème, curseur, options)  →  fichier prefs.json")
bullet("Journaux (logs)  →  dossier logs")

body(
    "Pour accéder à ce dossier directement : appuyez sur Windows + R, "
    "tapez  %LOCALAPPDATA%\\StuGoCO2  et appuyez sur Entrée.",
    before=6, after=4
)


# ── Désinstallation ────────────────────────────────────────────────────────────
section_title("Désinstaller le logiciel")

body(
    "Si vous souhaitez désinstaller StuGo CO2 Explorer, deux méthodes sont disponibles :",
    after=4
)
bullet(
    "Menu Démarrer : cherchez « StuGo » dans la liste des applications, "
    "faites un clic droit et choisissez Désinstaller."
)
bullet(
    "Paramètres Windows : allez dans Paramètres → Applications, "
    "cherchez « StuGo CO2 Explorer » dans la liste et cliquez sur Désinstaller."
)

body(
    "À la fin de la désinstallation, une fenêtre vous demandera si vous souhaitez "
    "également supprimer vos données personnelles (sessions, presets, journaux). "
    "Répondez Oui si vous voulez tout effacer, ou Non pour les garder.",
    before=6, after=4
)


# ── Problèmes fréquents ────────────────────────────────────────────────────────
section_title("Problèmes fréquents")

problems = [
    (
        "Le logiciel met longtemps à démarrer",
        "C'est normal au premier lancement (5 à 15 secondes). "
        "Les fois suivantes seront plus rapides."
    ),
    (
        "Windows affiche « Éditeur inconnu »",
        "Cliquez sur Exécuter quand même. Ce message apparaît pour tout logiciel "
        "sans certificat commercial — il n'indique pas que le fichier est dangereux."
    ),
    (
        "Je ne trouve pas le logiciel après l'installation",
        "Cherchez « StuGo » dans le Menu Démarrer. "
        "Si vous avez coché l'option raccourci, il est aussi sur votre Bureau."
    ),
    (
        "Erreur à l'ouverture d'un fichier Excel",
        "Vérifiez que le fichier correspond au format StuGo CO2 Explorer. "
        "Consultez l'onglet Aide dans le logiciel pour les détails."
    ),
    (
        "Le logiciel ne s'ouvre pas du tout",
        "Vérifiez que vous êtes bien sur Windows 10 ou Windows 11 en version 64 bits. "
        "Si le problème persiste, contactez la personne qui vous a fourni le logiciel."
    ),
]

for prob, sol in problems:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(prob)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    body(sol, before=0, after=5, indent=True)


# ── Pied de page discret ───────────────────────────────────────────────────────
doc.add_paragraph()
hrule()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_before = Pt(6)
rf = pf.add_run("StuGo CO2 Explorer v7.0  —  Document d'installation")
rf.font.name = 'Calibri'
rf.font.size = Pt(9)
rf.font.color.rgb = GRAY


doc.save('INSTALLATION.docx')
print('Document cree : INSTALLATION.docx')
