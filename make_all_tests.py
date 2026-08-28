"""
make_all_tests.py — Rapports de test StuGo CO2 Explorer
Génère : Rapport_Tests_v1_v3.docx / Rapport_Tests_v4_v5.docx / Rapport_Tests_v6.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE   = RGBColor(0x1A, 0x56, 0x9A)
GRAY   = RGBColor(0x55, 0x55, 0x55)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
ORANGE = RGBColor(0xC0, 0x70, 0x00)
RED    = RGBColor(0xBB, 0x22, 0x22)


def new_doc():
    doc = Document()
    doc.core_properties.author = "Remi Kalkan"
    doc.core_properties.last_modified_by = "Remi Kalkan"
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(2.8)
        s.right_margin  = Cm(2.2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc

def underline(p, color='1A569A'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    underline(p, 'CCCCCC')

def doc_title(doc, title, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(20)
    r.font.color.rgb = BLUE
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(14)
    r2 = p2.add_run(sub)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)
    r2.font.color.rgb = GRAY

def section_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(12)
    r.font.color.rgb = BLUE
    underline(p)

def sub_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11)
    r.font.color.rgb = BLACK

def body(doc, text, before=0, after=5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)

def kv(doc, key, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    rb = p.add_run(key + " : ")
    rb.bold = True; rb.font.name = 'Calibri'; rb.font.size = Pt(11)
    rn = p.add_run(val)
    rn.font.name = 'Calibri'; rn.font.size = Pt(11)

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.8)
    r = p.add_run("Note : " + text)
    r.italic = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
    r.font.color.rgb = GRAY

def doc_footer(doc, text):
    doc.add_paragraph()
    hrule(doc)
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(6)
    r = pf.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(9)
    r.font.color.rgb = GRAY

def add_test_table(doc, tests):
    """
    tests = list of (id, description, expected, actual, status)
    status: PASS | FAIL | PARTIEL | BLOQUÉ
    """
    STATUS_COLORS = {'PASS': GREEN, 'FAIL': RED, 'PARTIEL': ORANGE, 'BLOQUÉ': RED}
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    # widths
    for i, w in enumerate([Cm(1.2), Cm(3.8), Cm(3.5), Cm(4.2), Cm(1.6)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    # header
    hdr_texts = ["ID", "Description", "Attendu", "Obtenu", "Statut"]
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(hdr_texts):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(10)
        run.font.color.rgb = BLUE
    # rows
    for tid, desc, expected, actual, status in tests:
        row = tbl.add_row().cells
        color = STATUS_COLORS.get(status, GRAY)
        for col_i, (text, colored) in enumerate([
            (tid, False), (desc, False), (expected, False),
            (actual, status in ('FAIL', 'BLOQUÉ', 'PARTIEL')), (status, True)
        ]):
            p = row[col_i].paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9 if col_i in (2, 3) else 10)
            if colored:
                r.font.color.rgb = color
            if col_i in (0, 4):
                r.bold = True
    doc.add_paragraph()

def bug_entry(doc, bid, severity, title, repro, result, fix_version):
    colors = {'Critique': RED, 'Majeur': ORANGE, 'Mineur': GRAY}
    c = colors.get(severity, GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{bid}  [{severity}]  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(11); r1.font.color.rgb = c
    r2 = p.add_run(title)
    r2.bold = True; r2.font.name = 'Calibri'; r2.font.size = Pt(11)
    body(doc, f"Reproduction : {repro}",      before=0, after=1, indent=True)
    body(doc, f"Comportement constaté : {result}", before=0, after=1, indent=True)
    body(doc, f"Corrigé en : {fix_version}",  before=0, after=4, indent=True)


# =============================================================================
# DOC 1  Rapport de test v1 – v3
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Test  ·  Versions v1 · v2 · v3  ·  Janvier – Février 2026")
hrule(doc)

section_h(doc, "Informations générales")
kv(doc, "Versions testées",   "v1.0, v2.0, v3.0")
kv(doc, "Dates de test",      "16 jan. / 30 jan. / 20 fév. 2026  (en fin de chaque sprint)")
kv(doc, "Testeur",            "Équipe projet")
kv(doc, "Environnement",      "Windows 11 64 bits, Python 3.12, PyQt6, résolution 1920×1080")
kv(doc, "Méthode",            "Tests manuels — aucun test automatisé à ce stade")

section_h(doc, "Périmètre de test")
body(doc, "Les tests couvrent les fonctionnalités livrées à chaque sprint : import, tableau, graphiques et filtres. "
     "L'écran Comparaison est testé en mode ébauche uniquement (v3).")

# Table v1
sub_h(doc, "Tableau de cas de test — v1.0  (16 janvier 2026)")
add_test_table(doc, [
    ("TC-01", "Import d'un fichier .xlsx valide",
     "Données affichées dans le tableau",
     "Données affichées correctement",
     "PASS"),
    ("TC-02", "Import d'un fichier .xlsx avec colonne 'total_tco2' manquante",
     "Message d'erreur clair",
     "KeyError non géré — crash programme",
     "FAIL"),
    ("TC-03", "Import d'un fichier .xlsx ouvert dans Excel",
     "Message d'erreur 'fichier verrouillé'",
     "PermissionError non géré — crash programme",
     "FAIL"),
    ("TC-04", "Import d'un fichier vide (0 lignes de données)",
     "Message d'information 'aucune donnée'",
     "Tableau vide sans message — confusion utilisateur",
     "PARTIEL"),
    ("TC-05", "Redimensionnement de la fenêtre",
     "L'interface s'adapte proprement",
     "Tableau se redimensionne, boutons restent fixes",
     "PASS"),
    ("TC-06", "Import de deux fichiers successifs",
     "Les deux jeux de données fusionnés dans le tableau",
     "Non disponible en v1 — fonctionnalité absente",
     "BLOQUÉ"),
])

# Table v2
sub_h(doc, "Tableau de cas de test — v2.0  (30 janvier 2026)")
add_test_table(doc, [
    ("TC-07", "Affichage graphique barres verticales après import",
     "Graphique barres affiché, axe X = facultés",
     "Graphique affiché correctement",
     "PASS"),
    ("TC-08", "Affichage graphique camembert avec données normales",
     "Camembert affiché avec parts proportionnelles",
     "Affiché correctement si toutes les valeurs > 0",
     "PASS"),
    ("TC-09", "Affichage graphique camembert avec une faculté à 0 étudiant",
     "Graphique affiché sans cette faculté ou valeur ignorée",
     "ZeroDivisionError — crash programme",
     "FAIL"),
    ("TC-10", "Filtre par faculté : décocher une faculté",
     "Les données de cette faculté disparaissent du graphique",
     "Graphique mis à jour uniquement après clic Appliquer",
     "PARTIEL"),
    ("TC-11", "Filtre par zone CO2 : décocher zone 5",
     "Lignes zone 5 masquées en temps réel",
     "Mise à jour non temps réel — bouton Appliquer requis",
     "PARTIEL"),
    ("TC-12", "Compteur global étudiants et tCO2e",
     "Valeurs correctes affichées en haut",
     "Valeurs correctes",
     "PASS"),
    ("TC-13", "Labels de données avec 15+ facultés",
     "Labels lisibles, pas de chevauchement",
     "Labels se chevauchent — illisibles",
     "FAIL"),
])

# Table v3
sub_h(doc, "Tableau de cas de test — v3.0  (20 février 2026)")
add_test_table(doc, [
    ("TC-14", "Navigation sidebar vers chaque section",
     "Chaque section s'affiche au clic",
     "Navigation fonctionnelle",
     "PASS"),
    ("TC-15", "Indicateur section active dans la sidebar",
     "Le bouton de la section courante est visuellement actif",
     "Aucun état visuel actif — boutons tous identiques",
     "FAIL"),
    ("TC-16", "Graphique 3D barres verticales",
     "Barres 3D affichées avec rotation possible",
     "Graphique 3D affiché correctement",
     "PASS"),
    ("TC-17", "Graphique 3D type Treemap",
     "Message 'non supporté en 3D' ou fallback 2D",
     "Crash non géré — AttributeError mpl_toolkits",
     "FAIL"),
    ("TC-18", "Graphique 3D type Donut",
     "Message 'non supporté en 3D' ou fallback 2D",
     "Crash non géré — erreur matplotlib",
     "FAIL"),
    ("TC-19", "Écran Comparaison avec 3 fichiers",
     "3 graphiques côte à côte",
     "Graphiques affichés sans contrôles",
     "PARTIEL"),
    ("TC-20", "Interface sur écran 4K (DPI 200%)",
     "Interface lisible, tailles proportionnelles",
     "Boutons et textes minuscules — aucune adaptation DPI",
     "FAIL"),
    ("TC-21", "Page d'accueil statistiques",
     "6 compteurs mis à jour après import",
     "Compteurs corrects",
     "PASS"),
])

section_h(doc, "Rapport de bugs identifiés")
bug_entry(doc, "BUG-001", "Majeur",
          "Crash KeyError sur import fichier sans colonne 'total_tco2'",
          "Importer un fichier Excel ne contenant pas la colonne total_tco2",
          "KeyError non capturé — programme fermé sans message",
          "v4.0 — gestion d'exception ajoutée")
bug_entry(doc, "BUG-002", "Majeur",
          "ZeroDivisionError sur graphique camembert avec une valeur nulle",
          "Importer un fichier avec au moins une faculté à 0 étudiant, afficher camembert",
          "Crash ZeroDivisionError dans le calcul des parts",
          "v4.0 — guard ajouté avant calcul camembert")
bug_entry(doc, "BUG-003", "Majeur",
          "Filtres non temps réel",
          "Décocher une faculté dans le panneau de filtres",
          "Graphique non mis à jour — bouton Appliquer requis",
          "v3.0 — connexion directe aux signaux Qt")
bug_entry(doc, "BUG-004", "Majeur",
          "Crash en mode 3D sur les types Treemap et Donut",
          "Activer mode 3D puis sélectionner Treemap ou Donut",
          "AttributeError non géré — crash programme",
          "v4.0 — guard + fallback 2D ajoutés")
bug_entry(doc, "BUG-005", "Mineur",
          "Interface minuscule sur écrans HiDPI",
          "Lancer le logiciel sur un écran 4K avec DPI système à 150% ou 200%",
          "Tous les éléments UI sont minuscules — illisibles",
          "v4.0 — système de scaling DPI implémenté (shared/scaling.py)")

section_h(doc, "Synthèse")
body(doc, "Version v1.0 : 3 PASS / 2 FAIL / 1 PARTIEL / 1 BLOQUÉ — taux de succès : 43 %")
body(doc, "Version v2.0 : 3 PASS / 2 FAIL / 2 PARTIEL — taux de succès : 43 %")
body(doc, "Version v3.0 : 3 PASS / 4 FAIL / 1 PARTIEL — taux de succès : 37 %")
body(doc, "Bilan global v1→v3 : 9 cas réussis / 21 au total — taux de succès moyen : 43 %", before=6)
note(doc, "Le taux de succès faible est attendu à ce stade — les premières versions posent les bases "
     "sans finition. Les bugs identifiés ici sont tous résolus en v4 ou v5.")

doc_footer(doc, "StuGo CO2 Explorer  —  Rapport de Tests v1·v2·v3  ·  Janvier–Février 2026")
doc.save('Rapport_Tests_v1_v3.docx')
print('Rapport_Tests_v1_v3.docx créé')


# =============================================================================
# DOC 2  Rapport de test v4 – v5
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Test  ·  Versions v4 · v5  ·  Mars – Avril 2026")
hrule(doc)

section_h(doc, "Informations générales")
kv(doc, "Versions testées",   "v4.0, v5.0")
kv(doc, "Dates de test",      "20 mars / 25 avril 2026  (en fin de chaque sprint)")
kv(doc, "Testeur",            "Équipe projet")
kv(doc, "Environnement",      "Windows 11 64 bits, Python 3.12, écrans 1080p et 4K")
kv(doc, "Méthode",            "Tests manuels — aucun test automatisé à ce stade")

section_h(doc, "Périmètre de test")
body(doc, "Tests couvrant les nouvelles fonctionnalités : DPI scaling, thèmes, preset couleurs, "
     "export graphique, persistance session, page Aide et Admin Logs.")

# v4 table
sub_h(doc, "Tableau de cas de test — v4.0  (20 mars 2026)")
add_test_table(doc, [
    ("TC-22", "Interface sur écran 4K (DPI 200%)",
     "Tailles proportionnelles, lisible",
     "Interface correctement adaptée — DPI scaling fonctionnel",
     "PASS"),
    ("TC-23", "Application d'un thème prédéfini",
     "Toute l'interface change de couleurs instantanément",
     "Interface mise à jour correctement",
     "PASS"),
    ("TC-24", "Changer de thème plusieurs fois rapidement",
     "Chaque thème s'applique sans erreur",
     "Fonctionnel — aucune fuite mémoire observable",
     "PASS"),
    ("TC-25", "Preset personnalisé — clic couleur mode 18 clés",
     "Sélecteur de couleur s'ouvre, couleur mise à jour",
     "Fonctionne correctement en mode 18 clés",
     "PASS"),
    ("TC-26", "Preset personnalisé — clic couleur mode 6 clés",
     "Sélecteur de couleur s'ouvre, couleur mise à jour",
     "AttributeError : 'NoneType' object has no attribute 'setText' — crash",
     "FAIL"),
    ("TC-27", "Preset personnalisé — clic couleur mode 2 clés",
     "Sélecteur de couleur s'ouvre, couleur mise à jour",
     "AttributeError : 'NoneType' object has no attribute 'setText' — crash",
     "FAIL"),
    ("TC-28", "Sauvegarder un preset couleur en JSON",
     "Fichier .json créé dans %LOCALAPPDATA%\\StuGoCO2\\PresetColors",
     "Fichier créé correctement",
     "PASS"),
    ("TC-29", "Charger un preset couleur JSON au redémarrage",
     "Preset rechargé automatiquement",
     "Preset non rechargé — fichier ignoré au démarrage",
     "FAIL"),
    ("TC-30", "Graphique 3D Treemap",
     "Message 'non disponible en 3D' affiché, graphique en 2D",
     "Message affiché, fallback 2D fonctionnel",
     "PASS"),
    ("TC-31", "Sidebar : indicateur section active",
     "Section courante visuellement distinguée",
     "Bouton actif mis en surbrillance correctement",
     "PASS"),
])

# v5 table
sub_h(doc, "Tableau de cas de test — v5.0  (25 avril 2026)")
add_test_table(doc, [
    ("TC-32", "Export graphique en PNG",
     "Fichier PNG créé à l'emplacement choisi",
     "Fichier PNG créé correctement",
     "PASS"),
    ("TC-33", "Export graphique en PDF",
     "Fichier PDF créé avec graphique visible",
     "PDF créé, mais vide si graphique non encore rendu à l'écran",
     "PARTIEL"),
    ("TC-34", "Export graphique en SVG",
     "Fichier SVG créé, modifiable dans Inkscape",
     "SVG créé correctement",
     "PASS"),
    ("TC-35", "Export en lot depuis Comparaison (PNG)",
     "Fichiers PNG créés pour chaque graphique visible",
     "Tous les fichiers créés correctement",
     "PASS"),
    ("TC-36", "Sauvegarde de session à la fermeture",
     "Fichier session.json mis à jour",
     "Session sauvegardée correctement",
     "PASS"),
    ("TC-37", "Restauration de session au démarrage",
     "Fichiers Excel rechargés automatiquement",
     "Restauration OK si les fichiers sont toujours au même emplacement",
     "PASS"),
    ("TC-38", "Restauration session après déplacement des fichiers Excel",
     "Message d'erreur clair, session partiellement restaurée",
     "Échec silencieux — aucun message, données absentes",
     "FAIL"),
    ("TC-39", "Crash programme forcé (Gestionnaire des tâches) puis redémarrage",
     "Session partiellement restaurée avec avertissement",
     "Fichier session.json corrompu (JSON incomplet) — impossible à lire",
     "FAIL"),
    ("TC-40", "Page Aide : contenu du format de fichier",
     "Description des colonnes attendues affichée",
     "Contenu affiché correctement",
     "PASS"),
    ("TC-41", "Admin Logs : affichage des logs",
     "Liste des logs affichée, filtrable par date",
     "Logs affichés, filtres par date absents",
     "PARTIEL"),
    ("TC-42", "Preset couleurs — bug clic mode 6 clés (régression v4)",
     "Sélecteur de couleur s'ouvre normalement",
     "AttributeError toujours présent — bug non corrigé",
     "FAIL"),
])

section_h(doc, "Rapport de bugs identifiés")
bug_entry(doc, "BUG-006", "Critique",
          "Crash color picker mode 2/6 clés (régression présente depuis v4)",
          "Ouvrir Paramètres > Preset personnalisé, sélectionner mode 6 clés ou 2 clés, cliquer sur un bouton de couleur",
          "AttributeError: 'NoneType' object has no attribute 'setText' — crash programme",
          "v7.0 — tuple corrigé : btn, _ = self._color_btns[key]")
bug_entry(doc, "BUG-007", "Majeur",
          "Export PDF vide si graphique non visible",
          "Aller dans Graphique, ne pas attendre l'affichage, cliquer immédiatement Exporter PDF",
          "Fichier PDF créé mais vide — aucune figure exportée",
          "v5.0 — correction : forcer le rendu avant l'export")
bug_entry(doc, "BUG-008", "Majeur",
          "Restauration session silencieuse si fichier déplacé",
          "Importer un fichier, sauvegarder la session, déplacer le fichier Excel, redémarrer",
          "Session chargée sans message d'erreur, données absentes",
          "v6.0 — ajout d'un avertissement si fichier introuvable")
bug_entry(doc, "BUG-009", "Mineur",
          "Preset couleur JSON non rechargé au démarrage",
          "Créer un preset, le sauvegarder, relancer le logiciel",
          "Le preset n'apparaît pas dans la liste des thèmes — fichier ignoré",
          "v5.0 — scanning du dossier PresetColors au démarrage ajouté")

section_h(doc, "Synthèse")
body(doc, "Version v4.0 : 6 PASS / 3 FAIL / 1 PARTIEL — taux de succès : 60 %")
body(doc, "Version v5.0 : 6 PASS / 3 FAIL / 2 PARTIEL — taux de succès : 55 %")
body(doc, "Bilan global v4–v5 : 12 réussis / 21 au total — taux de succès moyen : 57 %", before=6)
note(doc, "Le bug color picker (TC-26, TC-27, TC-42) est le seul bug critique restant à la fin du sprint 5. "
     "Il sera résolu en v7.0.")

doc_footer(doc, "StuGo CO2 Explorer  —  Rapport de Tests v4·v5  ·  Mars–Avril 2026")
doc.save('Rapport_Tests_v4_v5.docx')
print('Rapport_Tests_v4_v5.docx créé')


# =============================================================================
# DOC 3  Rapport de test v6
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Test  ·  Version v6.0 → v6.1  ·  Avril – Mai 2026")
hrule(doc)

section_h(doc, "Informations générales")
kv(doc, "Versions testées",   "v6.0  (27 avr. 2026)  et  v6.1  (09 mai 2026)")
kv(doc, "Dates de test",      "30 avril 2026 (v6.0) / 12 mai 2026 (v6.1)")
kv(doc, "Testeur",            "Équipe projet")
kv(doc, "Environnement",      "Windows 11 64 bits, Python 3.12, PyInstaller 6.x, Inno Setup 6.3.3")
kv(doc, "Méthode",            "Tests manuels — aucun test automatisé")
kv(doc, "Périmètre",          "Splash screen, build EXE, installeur, régressions sur fonctionnalités v5")

section_h(doc, "Tableau de cas de test — v6.0  (30 avril 2026)")
add_test_table(doc, [
    ("TC-43", "Splash screen au démarrage",
     "Animation GIF affichée 3–5 secondes, puis fenêtre principale",
     "Splash screen fonctionnel",
     "PASS"),
    ("TC-44", "Build PyInstaller onefile — lancement de l'exe",
     "Programme démarre sans erreur",
     "ImportError pandas._libs.reduction — crash immédiat au démarrage",
     "FAIL"),
    ("TC-45", "Build PyInstaller onedir — lancement de l'exe",
     "Programme démarre normalement",
     "Programme démarre correctement en mode onedir",
     "PASS"),
    ("TC-46", "Installeur Inno Setup — installation sans droits admin",
     "Installation dans %LOCALAPPDATA%\\Programs\\StuGo CO2 Explorer",
     "Installation réussie dans le dossier personnel, pas de demande admin",
     "PASS"),
    ("TC-47", "Dossier d'installation après setup (mode onedir)",
     "Dossier propre avec 2 exécutables max",
     "Dossier _internal avec 200+ DLL visibles — non conforme",
     "FAIL"),
    ("TC-48", "Inno Setup : compilation du script .iss",
     "Setup.exe produit sans erreur",
     "Error 32 EndUpdateResource — Defender verrouille l'exe",
     "FAIL"),
    ("TC-49", "Preset personnalisé — mode 6 clés (régression color picker)",
     "Sélecteur de couleur fonctionne",
     "AttributeError NoneType.setText — crash (bug toujours présent)",
     "FAIL"),
    ("TC-50", "Import d'un fichier Excel valide (régression)",
     "Données affichées dans le tableau",
     "Import fonctionnel",
     "PASS"),
    ("TC-51", "Export PNG (régression)",
     "Fichier PNG créé",
     "Export fonctionnel",
     "PASS"),
    ("TC-52", "Numéro de version affiché dans la sidebar",
     "Affiche la version courante (v6.0 ou v6.1)",
     "Affiche 'v6.0' en v6.0 — correct",
     "PASS"),
    ("TC-53", "Numéro de version dans la page Aide",
     "Affiche 'v6.0' ou 'v6.1' cohérent avec la sidebar",
     "Affiche 'v6.1' en v6.0 — incohérence entre fichiers",
     "FAIL"),
    ("TC-54", "Import CSV StuGo",
     "Fichier CSV importé et données affichées",
     "Fonctionnalité absente — aucun support CSV",
     "BLOQUÉ"),
])

section_h(doc, "Tableau de cas de test — v6.1  (12 mai 2026)")
add_test_table(doc, [
    ("TC-55", "Build PyInstaller onefile après correction hiddenimports",
     "Exe démarre sans erreur",
     "ImportError pandas._libs.reduction toujours présent — non corrigé en v6.1",
     "FAIL"),
    ("TC-56", "Preset personnalisé — mode 6 clés après patch v6.1",
     "Sélecteur de couleur fonctionne",
     "Crash partiel réduit mais AttributeError persiste dans certains cas",
     "PARTIEL"),
    ("TC-57", "Installeur Inno Setup après exclusion Defender",
     "Setup.exe produit sans erreur",
     "Compilation réussie après ajout dist/ aux exclusions Defender",
     "PASS"),
    ("TC-58", "Cohérence des numéros de version entre tous les fichiers",
     "Même numéro dans sidebar, aide, paramètres, titre fenêtre",
     "Incohérence persistante : v6.0 et v6.1 mélangés selon les fichiers",
     "FAIL"),
    ("TC-59", "Import de 5 fichiers Excel simultanément (régression)",
     "5 jeux de données chargés et fusionnés",
     "Import fonctionnel",
     "PASS"),
    ("TC-60", "Comparaison côte à côte avec 4 fichiers sélectionnés",
     "4 graphiques en grille 2×2",
     "Affichage correct",
     "PASS"),
])

section_h(doc, "Rapport de bugs identifiés")
bug_entry(doc, "BUG-010", "Critique",
          "Build onefile crash au démarrage (ImportError pandas._libs.reduction)",
          "Builder avec PyInstaller en mode onefile, lancer l'exe produit",
          "ImportError: No module named 'pandas._libs.reduction' — programme ne démarre pas",
          "v7.0 — pandas._libs.reduction retiré des hiddenimports dans build_all.spec")
bug_entry(doc, "BUG-011", "Critique",
          "Crash color picker mode 2/6 clés — toujours présent en v6.1",
          "Paramètres > Preset personnalisé > mode 6 clés > clic sur un bouton de couleur",
          "AttributeError: 'NoneType' has no attribute 'setText' — crash programme",
          "v7.0 — btn, _ = self._color_btns[key]; appel lbl supprimé")
bug_entry(doc, "BUG-012", "Majeur",
          "Inno Setup Error 32 lors de la compilation",
          "Lancer ISCC.exe installer.iss immédiatement après un build PyInstaller",
          "Error 32: EndUpdateResource failed — Windows Defender scanne et verrouille l'exe",
          "Contournement v6.1 : ajouter dist/ aux exclusions Windows Defender")
bug_entry(doc, "BUG-013", "Majeur",
          "Références de version incohérentes entre fichiers",
          "Comparer le numéro de version affiché dans sidebar, aide, paramètres, titre fenêtre",
          "main_window.py: 'v6', sidebar: 'v6.0', help_page: 'v6.1', installer.iss: '6.1'",
          "v7.0 — toutes les références mises à jour à v7.0 de façon uniforme")
bug_entry(doc, "BUG-014", "Majeur",
          "Mode onedir : dossier _internal avec 200+ DLL visible après installation",
          "Installer le logiciel depuis StuGoCO2_Setup_v6.1.exe, ouvrir le dossier d'installation",
          "Dossier _internal contenant toutes les DLL est visible pour l'utilisateur",
          "v7.0 — passage en mode onefile : DLL embarquées dans l'exe, dossier _internal absent")

section_h(doc, "Synthèse")
body(doc, "Version v6.0 : 5 PASS / 6 FAIL / 1 BLOQUÉ — taux de succès : 38 %")
body(doc, "Version v6.1 : 3 PASS / 2 FAIL / 1 PARTIEL — taux de succès : 50 %")
body(doc, "Bilan global v6 : 8 réussis / 18 au total — taux de succès moyen : 44 %", before=6)
note(doc, "Le sprint 6 est le plus difficile du projet. Les deux bugs critiques (onefile et color picker) "
     "ne seront définitivement résolus qu'en v7.0. Le taux de réussite bas reflète "
     "la complexité des problèmes de distribution sous Windows.")

doc_footer(doc, "StuGo CO2 Explorer  —  Rapport de Tests v6.0·v6.1  ·  Avril–Mai 2026")
doc.save('Rapport_Tests_v6.docx')
print('Rapport_Tests_v6.docx créé')
