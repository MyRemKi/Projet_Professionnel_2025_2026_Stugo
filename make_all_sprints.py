"""
make_all_sprints.py — Rapports de sprint StuGo CO2 Explorer v1 → v7
Génère : Rapport_Sprint_v1_v2_v3.docx / Rapport_Sprint_v4_v5.docx
         Rapport_Sprint_v6.docx / Rapport_Projet_Sprint_v7.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE   = RGBColor(0x1A, 0x56, 0x9A)
GRAY   = RGBColor(0x55, 0x55, 0x55)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
ORANGE = RGBColor(0xC0, 0x70, 0x00)
RED    = RGBColor(0xBB, 0x22, 0x22)


def new_doc():
    doc = Document()
    doc.core_properties.author = "Remi Kalkan"
    doc.core_properties.last_modified_by = "Remi Kalkan"
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc

def underline(p, color='1A569A'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    underline(p, 'CCCCCC')

def doc_title(doc, title, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(22)
    r.font.color.rgb = BLUE
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(14)
    r2 = p2.add_run(sub)
    r2.font.name = 'Calibri'; r2.font.size = Pt(12)
    r2.font.color.rgb = GRAY

def sprint_title(doc, num, version, title, dates, pts=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"Sprint {num}  ·  {version}  —  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(14)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(title)
    r2.bold = True; r2.font.name = 'Calibri'; r2.font.size = Pt(14)
    r2.font.color.rgb = BLUE
    underline(p)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(8)
    suffix = f"  ·  Vélocité cible : {pts} pts" if pts else ""
    r3 = p2.add_run(dates + suffix)
    r3.italic = True; r3.font.name = 'Calibri'; r3.font.size = Pt(10)
    r3.font.color.rgb = GRAY

def section_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11)
    r.font.color.rgb = BLACK

def body(doc, text, before=0, after=5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)

def task(doc, label, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    rl = p.add_run(label + "  ")
    rl.bold = True; rl.font.name = 'Calibri'; rl.font.size = Pt(10)
    rl.font.color.rgb = color
    rt = p.add_run(text)
    rt.font.name = 'Calibri'; rt.font.size = Pt(11)

def done(doc, t): task(doc, "[✔ TERMINÉ]",  t, GREEN)
def partial(doc, t): task(doc, "[◑ PARTIEL]",    t, ORANGE)
def blocked(doc, t): task(doc, "[✖ BLOQUÉ]", t, RED)

def bug(doc, sev, text):
    c = {'Critique': RED, 'Majeur': ORANGE, 'Mineur': GRAY}.get(sev, GRAY)
    task(doc, f"[BUG {sev.upper()}]", text, c)

def rplus(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run("(+)  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.font.color.rgb = GREEN
    r2 = p.add_run(t); r2.font.name = 'Calibri'; r2.font.size = Pt(11)

def rminus(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run("(–)  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.font.color.rgb = RED
    r2 = p.add_run(t); r2.font.name = 'Calibri'; r2.font.size = Pt(11)

def action(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run("→  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(t); r2.font.name = 'Calibri'; r2.font.size = Pt(11)

def note(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.8)
    r = p.add_run("Note : " + t)
    r.italic = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
    r.font.color.rgb = GRAY

def chapter_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(14)
    r.font.color.rgb = BLUE
    underline(p)

def doc_footer(doc, text):
    doc.add_paragraph()
    hrule(doc)
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(6)
    r = pf.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(9)
    r.font.color.rgb = GRAY


# =============================================================================
# DOC 1  Rapport Sprint v1 · v2 · v3
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Sprint  ·  Versions v1 · v2 · v3  ·  Janvier – Février 2026")
hrule(doc)
body(doc,
     "Ce document regroupe les trois premiers sprints du projet StuGo CO2 Explorer. "
     "Il couvre la mise en place de la base technique, l’ajout des graphiques et filtres, "
     "et la restructuration en modules avec navigation et graphiques 3D.",
     before=10, after=8)

# Sprint 1
sprint_title(doc, "1", "v1.0", "Mise en place de la base technique",
             "05 janvier 2026 – 16 janvier 2026", "24")
section_h(doc, "Objectif du sprint")
body(doc, "Créer la première version fonctionnelle : fenêtre principale PyQt6, import d’un fichier Excel StuGo, affichage sous forme de tableau.")
section_h(doc, "Backlog et user stories")
bullet(doc, "[8 pts]  Créer la fenêtre principale PyQt6 avec layout de base")
bullet(doc, "[5 pts]  Lire un fichier .xlsx avec pandas et extraire les données")
bullet(doc, "[8 pts]  Afficher les données dans un QTableView")
bullet(doc, "[3 pts]  Bouton d’import connecté à QFileDialog")
section_h(doc, "État des tâches — fin de sprint")
done(doc,    "Fenêtre principale PyQt6 créée, redimensionnable, icône d’application")
done(doc,    "Lecture .xlsx avec pandas.read_excel, extraction des colonnes StuGo")
done(doc,    "QTableView fonctionnel avec les données importées")
done(doc,    "Bouton import connecté à QFileDialog (filtre .xlsx/.xls)")
partial(doc, "Validation des colonnes requises : présence vérifiée, message d’erreur absent")
blocked(doc, "Chargement multi-fichiers — prévu sprint 2")
section_h(doc, "Bugs et problèmes identifiés")
bug(doc, "Majeur",   "Crash KeyError si colonne 'total_tco2' absente — aucune gestion d’exception")
bug(doc, "Majeur",   "Crash si le fichier Excel est ouvert dans Excel au moment de l’import")
bug(doc, "Mineur",   "Interface sans style CSS — boutons et tableau visuellement bruts")
bug(doc, "Mineur",   "Aucun message si fichier vide ou corrompu")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 21 pts / 24.  Les 3 pts manquants (validation colonnes) reportés au sprint 2.")
section_h(doc, "Rétrospective")
rplus(doc,  "Premier build fonctionnel livré en 2 semaines comme prévu")
rplus(doc,  "PyQt6 + pandas identifiés comme stack stable et adapté")
rminus(doc, "Aucune gestion d’erreur — crash sur données inattendues")
rminus(doc, "Architecture monolithique (tout dans main.py) difficile à étendre")
action(doc, "Ajouter try/except sur toutes les opérations d’import — sprint 2")
action(doc, "Planifier refonte en modules presentation/, infrastructure/ — sprint 3")

# Sprint 2
sprint_title(doc, "2", "v2.0", "Graphiques 2D et filtres",
             "19 janvier 2026 – 30 janvier 2026", "29")
section_h(doc, "Objectif du sprint")
body(doc, "Intégrer matplotlib pour afficher des graphiques 2D et ajouter un panneau de filtres par faculté et par zone CO2.")
section_h(doc, "Backlog et user stories")
bullet(doc, "[8 pts]  Intégrer matplotlib dans un widget Qt (FigureCanvasQTAgg)")
bullet(doc, "[5 pts]  Graphique barres verticales — axe X : facultés, axe Y : tCO2e")
bullet(doc, "[3 pts]  Graphique camembert")
bullet(doc, "[5 pts]  Filtre par faculté (cases à cocher)")
bullet(doc, "[5 pts]  Filtre par zone CO2 (1 à 5)")
bullet(doc, "[3 pts]  Compteur global étudiants et tCO2e total affiché en haut")
section_h(doc, "État des tâches — fin de sprint")
done(doc,    "matplotlib intégré via FigureCanvasQTAgg dans un QWidget")
done(doc,    "Graphiques barres verticales et horizontales")
done(doc,    "Filtre par faculté fonctionnel (mise à jour via bouton Appliquer)")
done(doc,    "Compteur global affiché")
partial(doc, "Filtre par zone CO2 : cases à cocher visibles, mise à jour non temps réel")
partial(doc, "Graphique camembert : affiché mais crash si une valeur est nulle (ZeroDivisionError)")
blocked(doc, "Mise à jour temps réel des filtres (signal/slot complexe) — reporté sprint 3")
section_h(doc, "Bugs et problèmes identifiés")
bug(doc, "Majeur", "Camembert : ZeroDivisionError si une part vaut 0 — pas de guard")
bug(doc, "Majeur", "Filtre par zone ne se met pas à jour sans cliquer Appliquer — UX dégradé")
bug(doc, "Mineur", "Labels de données sur barres se chevauchent avec beaucoup de facultés")
bug(doc, "Mineur", "Pas de type Donut ni Treemap — uniquement barres et camembert")
bug(doc, "Mineur", "Figure matplotlib non redimensionnée quand la fenêtre change de taille")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 24 pts / 29.  Donut et Treemap non livrés, reportés sprint 3.")
section_h(doc, "Rétrospective")
rplus(doc,  "Graphiques fonctionnels dès la première itération")
rplus(doc,  "L’interface est utilisable pour des démos simples")
rminus(doc, "Filtres non temps réel — pénalise l’expérience utilisateur")
rminus(doc, "Pas de gestion des cas limites (valeur 0, colonne vide)")
action(doc, "Connecter les filtres directement aux signaux Qt — sprint 3")
action(doc, "Ajouter guard avant tout calcul de graphique — sprint 3")

# Sprint 3
sprint_title(doc, "3", "v3.0", "Navigation, graphiques 3D, architecture modulaire",
             "02 février 2026 – 20 février 2026", "42")
section_h(doc, "Objectif du sprint")
body(doc, "Restructurer l’application en modules, introduire la sidebar de navigation, les graphiques 3D et l’ébauche de l’écran Comparaison.")
section_h(doc, "Backlog et user stories")
bullet(doc, "[8 pts]  Sidebar de navigation latérale (QWidget vertical avec boutons)")
bullet(doc, "[5 pts]  Page d’accueil avec statistiques globales")
bullet(doc, "[13 pts] Graphiques 3D avec mpl_toolkits.mplot3d")
bullet(doc, "[8 pts]  Écran Comparaison (ébauche — liste fichiers + graphique)")
bullet(doc, "[8 pts]  Refonte en modules : presentation/, infrastructure/, domain/")
section_h(doc, "État des tâches — fin de sprint")
done(doc,    "Sidebar de navigation avec 6 sections (Accueil, Import, Tableau, Graphique, Comparaison, Paramètres)")
done(doc,    "Page d’accueil avec 6 compteurs statistiques")
done(doc,    "Graphiques 3D barres et camémbert fonctionnels")
done(doc,    "Architecture modulaire mise en place")
partial(doc, "Écran Comparaison : liste + graphique mais aucun contrôle de tri ni sélection multiple")
partial(doc, "Filtres par zone maintenant en temps réel — filtres avancés (min/max) manquants")
blocked(doc, "Treemap 3D non réalisé — matplotlib ne supporte pas le rendu 3D pour ce type")
section_h(doc, "Bugs et problèmes identifiés")
bug(doc, "Majeur", "Crash non géré en mode 3D avec Treemap et Donut — pas de fallback 2D")
bug(doc, "Majeur", "Sidebar ne montre pas la section active (aucun état visuel sélectionné)")
bug(doc, "Mineur", "Comparaison instable avec plus de 5 fichiers chargés simultanément")
bug(doc, "Mineur", "Pas de DPI awareness : UI minuscule sur écrans 4K / HiDPI")
bug(doc, "Mineur", "Rotation 3D uniquement via souris, pas de curseurs UI Elévation/Azimut")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 37 pts / 42.  Sprint allongé de 3 jours pour finir la modularisation.")
section_h(doc, "Rétrospective")
rplus(doc,  "Architecture modulaire posée — plus facile à maintenir et étendre")
rplus(doc,  "Graphiques 3D appréciés visuellement lors des démos")
rminus(doc, "DPI awareness complètement absent — bloquant pour utilisateurs HiDPI")
rminus(doc, "Sprint trop large (42 pts / 2 semaines) — estimation trop optimiste")
rminus(doc, "Trop de bugs bloquants laissés sans traitement — dette technique croissante")
action(doc, "Sprint 4 dédié : DPI awareness + thèmes — ne pas ajouter de features")
action(doc, "Écrire guards de type avant chaque rendu 3D")

doc_footer(doc, "StuGo CO2 Explorer  —  Rapports de Sprint v1 · v2 · v3  ·  Janvier–Février 2026")
doc.save('Rapport_Sprint_v1_v2_v3.docx')
print('Rapport_Sprint_v1_v2_v3.docx créé')


# =============================================================================
# DOC 2  Rapport Sprint v4 · v5
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Sprint  ·  Versions v4 · v5  ·  Février – Avril 2026")
hrule(doc)
body(doc,
     "Ce document couvre les sprints 4 et 5 : scaling DPI, thèmes de couleurs, "
     "page Paramètres, persistance de session, export graphique et outil Admin Logs.",
     before=10, after=8)

# Sprint 4
sprint_title(doc, "4", "v4.0", "DPI Scaling, Thèmes, Paramètres",
             "23 février 2026 – 20 mars 2026", "39")
section_h(doc, "Objectif du sprint")
body(doc, "Rendre l’interface adaptative à tous les écrans via un système de scaling DPI, ajouter 7 thèmes visuels prédéfinis et une page Paramètres complète.")
section_h(doc, "Backlog et user stories")
bullet(doc, "[13 pts] Système de scaling DPI centralisé (shared/scaling.py — classe S)")
bullet(doc, "[8 pts]  7 thèmes prédéfinis avec application en temps réel")
bullet(doc, "[8 pts]  Page Paramètres avec sélecteur de thème et de curseur")
bullet(doc, "[8 pts]  Preset de couleurs personnalisé (3 modes : 2 / 6 / 18 clés)")
bullet(doc, "[5 pts]  Sélecteur de curseur de souris (8 styles)")
section_h(doc, "État des tâches — fin de sprint")
done(doc,    "Système DPI scaling S fonctionnel : toutes les tailles adaptées automatiquement")
done(doc,    "7 thèmes (Arctic Ink, Sombre, Clair Pro, Violet Nuit, Bleu Marine, Solarized, Sable Clair)")
done(doc,    "Page Paramètres avec sélecteur de thème et de curseur")
done(doc,    "Preset personnalisé : affichage 18 boutons couleur, modes 2/6/18 clés")
partial(doc, "Sauvegarde preset en JSON : fichier créé mais rechargement non fonctionnel au démarrage")
partial(doc, "Color picker mode 2/6 clés : crash AttributeError lbl.setText (NoneType) — non corrigé")
blocked(doc, "Application du preset aux graphiques matplotlib — non prioritaire, reporté")
section_h(doc, "Bugs et problèmes identifiés")
bug(doc, "Critique", "Crash AttributeError dans settings_page.py : lbl.setText() appelé sur None. "
         "Le tuple _color_btns stocke (btn, None) mais le code suppose (btn, lbl). "
         "Présent en mode 2 et 6 clés. Reporté sprint 5.")
bug(doc, "Majeur",   "Preset sauvegardé en JSON mais non rechargé au démarrage — persistance rompue")
bug(doc, "Majeur",   "Changement de thème ne s’applique pas aux couleurs des graphiques matplotlib")
bug(doc, "Mineur",   "Label version sidebar hardcodé à 'v1.0' — pas mis à jour")
bug(doc, "Mineur",   "Bouton Aperçu n’annule pas automatiquement si on change de thème")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 34 pts / 39.  Bug critique color picker non résolu — reporté.")
section_h(doc, "Rétrospective")
rplus(doc,  "DPI scaling résout complètement les problèmes d’affichage HiDPI")
rplus(doc,  "Les 7 thèmes donnent un aspect professionnel au logiciel")
rminus(doc, "Bug critique color picker non corrigé — accumulation de dette technique")
rminus(doc, "Persistance preset cassée — fonctionnalité livrée à moitié")
action(doc, "Prioriser le bug color picker en sprint 5 ou 6 au plus tard")
action(doc, "Tester la sauvegarde/restauration preset en isolation avant intégration")

# Sprint 5
sprint_title(doc, "5", "v5.0", "Export, Session, Aide, Admin Logs",
             "23 mars 2026 – 25 avril 2026", "37")
section_h(doc, "Objectif du sprint")
body(doc, "Ajouter l’export de graphiques (PNG/PDF/SVG), la persistance de session entre les démarrages, une page d’aide intégrée et l’outil Admin Logs.")
section_h(doc, "Backlog et user stories")
bullet(doc, "[8 pts]  Export graphique PNG / PDF / SVG")
bullet(doc, "[8 pts]  Sauvegarde et restauration de session (JSON dans %LOCALAPPDATA%)")
bullet(doc, "[5 pts]  Page Aide avec description du format de fichier attendu")
bullet(doc, "[8 pts]  Admin Logs viewer — lecture et affichage des fichiers de log")
bullet(doc, "[3 pts]  Label de version dans la sidebar")
bullet(doc, "[5 pts]  Filtres avancés : min/max étudiants, min/max tCO2e, masquer zéros")
section_h(doc, "État des tâches — fin de sprint")
done(doc,    "Export PNG / PDF / SVG depuis l’écran Graphique")
done(doc,    "Export en lot (tous graphiques Comparaison) dans un dossier")
done(doc,    "Sauvegarde session JSON dans %LOCALAPPDATA%\\StuGoCO2")
done(doc,    "Page Aide avec format de fichier")
done(doc,    "Admin Logs viewer (lecture des fichiers .log)")
done(doc,    "Filtres avancés min/max étudiants et tCO2e")
partial(doc, "Label version sidebar : affiché mais valeur hardcodée 'v6.0' (pas dynamique)")
partial(doc, "Restauration session : fonctionne si les fichiers n’ont pas bougé, échoue sinon")
blocked(doc, "Admin Logs : filtres par date non implémentés")
section_h(doc, "Bugs et problèmes identifiés")
bug(doc, "Critique", "Bug color picker mode 2/6 clés toujours présent depuis sprint 4 — non résolu")
bug(doc, "Majeur",   "Restauration session silencieusement échoue si les fichiers Excel ont été déplacés")
bug(doc, "Majeur",   "Export PDF produit un fichier vide si le graphique n’a pas encore été rendu à l’écran")
bug(doc, "Mineur",   "Session corrompue si le programme est forcé à quitter (JSON partiel)")
bug(doc, "Mineur",   "Label version sidebar hardcodé à 'v6.0' même en v5 — source de confusion")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 34 pts / 37.  Sprint allongé de 2 semaines — durée réelle : 5 semaines au lieu de 3.")
section_h(doc, "Rétrospective")
rplus(doc,  "Export PNG/PDF/SVG : fonctionnalité clé livrée et validée")
rplus(doc,  "Session persistence : très utile pour les reprises de travail")
rminus(doc, "Sprint allongé +2 semaines — premier retard significatif sur le planning")
rminus(doc, "Bug color picker reporté une 2ème fois — mauvaise impression en démo")
rminus(doc, "Références de version incohérentes entre fichiers (v1, v5, v6.0...)")
action(doc, "Sprint 6 : corriger le bug color picker EN PREMIER, avant toute nouvelle feature")
action(doc, "Mettre en place une convention de versioning unique dans tous les fichiers")

doc_footer(doc, "StuGo CO2 Explorer  —  Rapports de Sprint v4 · v5  ·  Février–Avril 2026")
doc.save('Rapport_Sprint_v4_v5.docx')
print('Rapport_Sprint_v4_v5.docx créé')


# =============================================================================
# DOC 3  Rapport Sprint v6
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Sprint  ·  Version v6.0 → v6.1  ·  Avril – Mai 2026")
hrule(doc)
body(doc,
     "Ce document couvre le sprint 6, centré sur la livraison d’un exécutable distributable "
     "(splash screen, build PyInstaller, installeur Inno Setup). Ce sprint a connu des retards "
     "significatifs et plusieurs bugs non résolus, notamment le crash color picker reporté depuis le sprint 4.",
     before=10, after=8)

sprint_title(doc, "6", "v6.0 → v6.1", "Build EXE, Splash Screen, Installeur",
             "27 avril 2026 – 23 mai 2026", "47")
section_h(doc, "Objectif du sprint")
body(doc, "Produire un exécutable unique installable sur Windows, avec splash screen animé et installeur Inno Setup. Corriger le bug color picker. Préparer la livraison.")
section_h(doc, "Backlog et user stories")
bullet(doc, "[8 pts]  Splash screen animé au démarrage (GIF + barre de chargement)")
bullet(doc, "[13 pts] Build PyInstaller onefile — un seul .exe sans dossier _internal")
bullet(doc, "[8 pts]  Script Inno Setup pour installeur Windows (sans droits admin)")
bullet(doc, "[5 pts]  Corriger le bug crash color picker (reporté depuis sprint 4)")
bullet(doc, "[5 pts]  Unifier les références de version dans tous les fichiers")
bullet(doc, "[8 pts]  Import de fichiers CSV StuGo")
section_h(doc, "État des tâches — fin de sprint")
done(doc,    "Splash screen animé fonctionnel (GIF + texte de chargement)")
done(doc,    "Installeur Inno Setup v6.1 fonctionnel — installation sans droits admin confirmée")
partial(doc, "Build PyInstaller : fonctionnel en mode onedir (dossier _internal visible) — onefile échoue")
partial(doc, "Bug color picker : identifié (lbl.setText sur NoneType) mais correction incomplète dans v6.0 "
         "— v6.1 publiée en urgence sans résoudre complètement")
blocked(doc, "Build onefile : ImportError pandas._libs.reduction — module supprimé dans pandas >= 2.0")
blocked(doc, "Unification des références de version : identifiée mais non faite — certains fichiers v6.0, d’autres v6.1")
blocked(doc, "Import CSV : non commencé — manque de temps")
section_h(doc, "Bugs et problèmes identifiés")
bug(doc, "Critique", "Build onefile crash au démarrage — ImportError pandas._libs.reduction "
         "(module n’existe plus dans pandas >= 2.0). Solution : retirer cette entrée des hiddenimports.")
bug(doc, "Critique", "Crash AttributeError settings_page.py : lbl.setText(c.name()) appelé sur None. "
         "Tuple _color_btns[key] = (btn, None) mais code suppose (btn, lbl). "
         "Présent depuis sprint 4 — toujours non résolu.")
bug(doc, "Majeur",   "Références de version incohérentes : 'v6.0' dans nav_sidebar.py et settings_page.py, "
         "'v6.1' dans help_page.py et installer.iss, 'v6' dans main_window.py.")
bug(doc, "Majeur",   "Mode onedir laisse un dossier _internal avec 200+ fichiers DLL visible à l’installation")
bug(doc, "Majeur",   "Inno Setup Error 32 (EndUpdateResource) : Windows Defender verrouille le nouvel exe. "
         "Contournement : ajouter dist/ aux exclusions Defender.")
bug(doc, "Mineur",   "Import CSV absent — fonctionnalité demandée non livrée")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 21 pts / 47.  Durée prévue : 2 semaines. Durée réelle : 4 semaines. "
     "Retard cumulé depuis mai : 3 semaines.")
section_h(doc, "Rétrospective")
rplus(doc,  "Installeur Inno Setup livré et fonctionnel — installation sans droits admin validée")
rplus(doc,  "Splash screen améliore l’expérience au premier démarrage")
rminus(doc, "Build onefile échoué — objectif principal du sprint non atteint")
rminus(doc, "Bug color picker reporté pour la 3ème fois (sprints 4, 5, 6) — problème de priorisation")
rminus(doc, "Version incohérente entre fichiers — absence de constante centrale dès le début")
rminus(doc, "Retard de 4 semaines sur planning initial — livraison prévue fin avril, effective fin mai")
action(doc, "Sprint 7 : bug color picker EN PREMIER (5 pts, estimé 2h max)")
action(doc, "Sprint 7 : résoudre build onefile (fix hiddenimports pandas._libs.reduction)")
action(doc, "Sprint 7 : unifier TOUTES les versions à v7.0 dans tous les fichiers")
action(doc, "Sprint 7 : implémenter l’import CSV")
action(doc, "Ne pas ajouter de nouvelles features au sprint 7 — stabilisation et livraison uniquement")
note(doc, "La version v6.1 a été publiée en cours de sprint pour corriger partiellement le crash, "
     "mais le bug n’a été complètement résolu qu’en v7.0.")

doc_footer(doc, "StuGo CO2 Explorer  —  Rapport de Sprint v6  ·  Avril–Mai 2026")
doc.save('Rapport_Sprint_v6.docx')
print('Rapport_Sprint_v6.docx créé')


# =============================================================================
# DOC 4  Rapport de Projet + Sprint v7
# =============================================================================
doc = new_doc()
doc_title(doc, "StuGo CO2 Explorer",
          "Rapport de Projet & Sprint Final  ·  Version v7.0  ·  Mai – Juin 2026")
hrule(doc)
body(doc,
     "Ce document est le rapport de clôture du projet StuGo CO2 Explorer. "
     "Il présente l’historique complet du projet, le bilan du sprint final (v7.0), "
     "le récapitulatif de tous les problèmes rencontrés et leurs solutions, "
     "ainsi qu’une analyse du retard accumulé depuis le lancement en janvier 2026.",
     before=10, after=8)

chapter_h(doc, "0.  Présentation du projet")
body(doc, "StuGo CO2 Explorer est un logiciel de bureau développé pour analyser et visualiser "
     "les émissions de CO2 liées aux déplacements d’étudiants en mobilité internationale. "
     "Il s’adresse aux équipes de relations internationales des établissements d’enseignement "
     "supérieur qui gèrent des données de mobilité au format StuGo (fichiers Excel structurés).", after=6)
section_h(doc, "Stack technique")
bullet(doc, "Langage : Python 3.12")
bullet(doc, "Interface graphique : PyQt6")
bullet(doc, "Traitement des données : pandas, numpy")
bullet(doc, "Graphiques : matplotlib 2D + 3D (mpl_toolkits)")
bullet(doc, "Distribution : PyInstaller 6.x (onefile) + Inno Setup 6")
bullet(doc, "Plateforme cible : Windows 10 / 11, 64 bits")
section_h(doc, "Architecture des modules")
bullet(doc, "shared/          — constantes, scaling DPI, chemins, utilitaires couleurs")
bullet(doc, "domain/          — value objects, interfaces repositories")
bullet(doc, "infrastructure/  — extracteurs Excel/CSV, modèles pandas, persistance")
bullet(doc, "app/             — services, commandes, bus d’événements, bootstrap")
bullet(doc, "presentation/    — fenêtre, pages, sidebar, thèmes, widgets")
bullet(doc, "rendering/       — stratégies de rendu 2D/3D, factory")

chapter_h(doc, "1.  Historique des versions")
versions_list = [
    ("v1.0",     "05–16 jan. 2026",    "Base : fenêtre PyQt6, import Excel, tableau de données"),
    ("v2.0",     "19–30 jan. 2026",    "Graphiques 2D (barres, camémbert), filtres faculté/zone"),
    ("v3.0",     "02–20 fév. 2026","Navigation sidebar, graphiques 3D, architecture modulaire"),
    ("v4.0",     "23 fév.–20 mar.","Scaling DPI, 7 thèmes, Paramètres, preset couleurs — bug color picker introduit"),
    ("v5.0",     "23 mar.–25 avr.","Export PNG/PDF/SVG, session, Aide, Admin Logs — bug color picker persistant"),
    ("v6.0/6.1", "27 avr.–23 mai","Splash screen, installeur Inno Setup, onefile échoué, bug partiellement corrigé"),
    ("v7.0",     "26 mai–06 jun.","Build onefile, fix crash color picker, import CSV, version unifiée, livraison finale"),
]
for ver, dates, desc in versions_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f"{ver:<12}")
    r1.bold = True; r1.font.name = 'Courier New'; r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(f"({dates})  —  {desc}")
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)

chapter_h(doc, "2.  Sprint 7  ·  v7.0  —  Stabilisation et livraison")
p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(8)
r3 = p2.add_run("26 mai 2026 – 06 juin 2026  ·  Vélocité cible : 36 pts")
r3.italic = True; r3.font.name = 'Calibri'; r3.font.size = Pt(10); r3.font.color.rgb = GRAY

section_h(doc, "Objectif du sprint")
body(doc, "Sprint de clôture. Aucune nouvelle fonctionnalité sauf l’import CSV planifié. "
     "Priorité : stabiliser, corriger les bugs critiques en suspens, livrer un exécutable "
     "onefile propre et une documentation à jour.")
section_h(doc, "Backlog")
bullet(doc, "[5 pts]  BUGFIX : crash color picker settings_page.py (AttributeError lbl.setText sur NoneType)")
bullet(doc, "[13 pts] Import de fichiers CSV exportés par StuGo CO2 Explorer")
bullet(doc, "[8 pts]  Build PyInstaller onefile (fix hiddenimports pandas._libs.reduction)")
bullet(doc, "[5 pts]  Unifier toutes les références de version à v7.0")
bullet(doc, "[3 pts]  Produire StuGoCO2_Setup_v7.0.exe sans écraser StuGoCO2_Setup_v6.1.exe")
bullet(doc, "[2 pts]  Régénérer INSTALLATION.docx et GUIDE_UTILISATEUR.docx en v7.0")
section_h(doc, "État des tâches — fin de sprint")
done(doc, "Bug color picker corrigé : btn, _ = self._color_btns[key] — appel lbl.setText supprimé")
done(doc, "csv_extractor.py : validation colonnes requises, parsing virgule décimale, groupement (fichier, sheet_id)")
done(doc, "extractor_factory.py : routing .xlsx/.xls → excel_extractor, .csv → csv_extractor")
done(doc, "build_all.spec : pandas._libs.reduction retiré des hiddenimports — build onefile OK")
done(doc, "main_window.py, nav_sidebar.py, settings_page.py, help_page.py — tous mis à jour à v7.0")
done(doc, "installer.iss : AppVersion 7.0, OutputBaseFilename StuGoCO2_Setup_v7.0")
done(doc, "INSTALLATION.docx et GUIDE_UTILISATEUR.docx régénérés en v7.0")
done(doc, "StuGoCO2_Setup_v7.0.exe produit (101 Mo) — v6.1 non écrasé (440 Mo conservé)")
section_h(doc, "Vélocité")
body(doc, "Réalisé : 36 pts / 36.  Sprint le plus court du projet (10 jours ouvrables). Toutes les tâches livrées.")

chapter_h(doc, "3.  Récapitulatif des problèmes et solutions")
problems = [
    ("Crash color picker (AttributeError : NoneType.setText)",
     "Sprints 4 à 6 — 3 sprints non résolu",
     "settings_page.py : _color_btns[key] retournait (btn, None) mais le code appelait lbl.setText(). "
     "Corrigé en v7.0 : btn, _ = self._color_btns[key]; appel lbl supprimé."),
    ("Build onefile PyInstaller échoue (ImportError pandas._libs.reduction)",
     "Sprint 6 — 1 sprint bloqué",
     "pandas._libs.reduction n’existe plus dans pandas >= 2.0. "
     "Retiré des hiddenimports dans build_all.spec. Build onefile opérationnel en v7.0."),
    ("Inno Setup Error 32 (EndUpdateResource)",
     "Sprint 6",
     "Windows Defender scannait et verrouillait le nouvel exe avant qu’Inno Setup puisse le signer. "
     "Solution : ajouter le dossier dist/ aux exclusions Windows Defender."),
    ("Références de version incohérentes entre fichiers",
     "Sprints 5 à 6 — 2 sprints",
     "Chaque fichier avait sa propre valeur hardcodée (v6.0, v6.1, v6, v1...). "
     "Corrigé en v7.0 : grep systématique sur tous les .py et .iss, mise à jour uniforme v7.0."),
    ("Build PyInstaller en mode onedir (dossier _internal visible)",
     "Sprint 6",
     "Mode onedir laissait 200+ DLL visibles dans le dossier d’installation. "
     "Résolu en passant au mode onefile : tous les fichiers embarqués dans l’exe."),
    ("Import CSV non disponible",
     "Sprints 5 à 6 — 2 sprints reporté",
     "Fonctionnalité demandée mais jamais priorisée. "
     "Implémentée en v7.0 via csv_extractor.py avec validation des 8 colonnes requises."),
    ("Filtres non temps réel",
     "Sprint 2",
     "Nécessitaient un bouton Appliquer. Corrigé en sprint 3 en connectant "
     "les filtres directement aux signaux Qt."),
    ("Crash graphique 3D sur types non supportés (Treemap, Donut)",
     "Sprint 3",
     "Pas de guard avant les appels 3D. Corrigé en sprint 4 avec message d’information "
     "et fallback automatique en 2D."),
    ("Interface minuscule sur écrans HiDPI/4K",
     "Sprints 1 à 3 — 3 sprints",
     "Aucun DPI awareness. Corrigé en sprint 4 avec shared/scaling.py : "
     "toutes les tailles calculées à partir de S.dpi_scale."),
]
for title_p, duration, solution in problems:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(title_p)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(1)
    p2.paragraph_format.left_indent  = Cm(0.5)
    r2 = p2.add_run(f"Durée : {duration}")
    r2.italic = True; r2.font.name = 'Calibri'; r2.font.size = Pt(10); r2.font.color.rgb = ORANGE
    body(doc, solution, before=1, after=4, indent=True)

chapter_h(doc, "4.  Analyse du retard et bilan planning")
body(doc, "Le projet a été initié en janvier 2026 avec un planning initial de 3 mois "
     "(janvier – fin mars). La livraison finale a eu lieu début juin 2026, "
     "soit un retard de 2 mois et demi.", after=6)
section_h(doc, "Chronologie prévue vs réelle")
planning = [
    ("v1 – v3", "Jan – Fév 2026", "05 jan. – 20 fév.",  "Conforme"),
    ("v4",       "Début mars 2026",   "23 fév. – 20 mar.","Conforme"),
    ("v5",       "Mars 2026",          "23 mar. – 25 avr.","Retard +3 sem."),
    ("v6",       "Fin mars 2026",      "27 avr. – 23 mai", "Retard +7 sem."),
    ("v7",       "Fin avril 2026",     "26 mai – 06 jun.", "Retard +6 sem."),
]
for ver, prev, real, delta in planning:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{ver:<10}")
    r1.font.name = 'Courier New'; r1.font.size = Pt(10)
    r2 = p.add_run(f"  Prévu : {prev:<22}  Réel : {real:<24}  {delta}")
    r2.font.name = 'Calibri'; r2.font.size = Pt(10)
    if "Retard" in delta:
        r2.font.color.rgb = ORANGE

section_h(doc, "Causes principales du retard")
bullet(doc, "Bug color picker non priorisé pendant 3 sprints — coût d’opportunité élevé")
bullet(doc, "Build onefile PyInstaller sous-estimé (incompatibilités pandas, Defender) — prévu 1 semaine, réalisé 3 semaines")
bullet(doc, "Sprints trop chargés — estimation trop optimiste (v5 : 37 pts pour 3 semaines)")
bullet(doc, "Absence de tests automatisés — chaque correction pouvait régresser les sprints précédents")
bullet(doc, "Découverte tardive des contraintes Windows Defender (Error 32 Inno Setup)")

section_h(doc, "Rétrospective globale du projet")
rplus(doc, "Toutes les fonctionnalités clés livrées : import Excel/CSV, tableau, graphiques 2D/3D, comparaison, export, thèmes, session")
rplus(doc, "Architecture modulaire propre et extensible (6 couches séparées)")
rplus(doc, "Exécutable onefile installable sans droits admin — intégration Windows native")
rplus(doc, "Documentation complète : procédure d’installation + guide utilisateur")
rminus(doc, "Retard de 2 mois et demi sur le planning initial")
rminus(doc, "Bug critique laissé 3 sprints sans correction — manque de discipline de priorisation")
rminus(doc, "Aucun test automatisé — tout testé manuellement, source de régressions")
rminus(doc, "Versioning incohérent pendant 2 sprints — absence de constante centrale")
action(doc, "Projets futurs : définir la convention de versioning dès le sprint 1 (constante centrale)")
action(doc, "Projets futurs : tester le build distributable dès le sprint 2 pour détecter les incompatibilités tôt")
action(doc, "Projets futurs : ne jamais reporter un bug critique plus d’un sprint")
action(doc, "Projets futurs : mettre en place des tests unitaires dès le sprint 3 (extracteurs, services)")

doc_footer(doc, "StuGo CO2 Explorer v7.0  —  Rapport de Projet & Sprint Final  ·  Juin 2026")
doc.save('Rapport_Projet_Sprint_v7.docx')
print('Rapport_Projet_Sprint_v7.docx créé')
