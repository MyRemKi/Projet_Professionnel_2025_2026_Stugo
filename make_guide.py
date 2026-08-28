"""
Génère GUIDE_UTILISATEUR.docx — guide d'utilisation StuGo CO2 Explorer v7.0
Style : sobre, lisible, humain. Pas de boîtes colorées, pas de mise en page bling.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


doc = Document()
doc.core_properties.author = "Remi Kalkan"
doc.core_properties.last_modified_by = "Remi Kalkan"

for s in doc.sections:
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

BLUE  = RGBColor(0x1A, 0x56, 0x9A)
GRAY  = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)


# ── Helpers ──────────────────────────────────────────────────────────────────

def underline_para(p, color='1A569A'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def hrule(light=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    underline_para(p, 'CCCCCC' if light else 'AAAAAA')

def title_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(24)
    r.font.color.rgb = BLUE

def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(20)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(12)
    r.font.color.rgb = GRAY

def chapter(num, text):
    """Titre de chapitre principal (ex: 1. Import des fichiers)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(8)
    r1 = p.add_run(f"{num}.  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(15)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.bold = True; r2.font.name = 'Calibri'; r2.font.size = Pt(15)
    r2.font.color.rgb = BLUE
    underline_para(p)

def section(text):
    """Sous-titre de section dans un chapitre"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(12)
    r.font.color.rgb = BLACK

def body(text, before=0, after=5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.5)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.8)
    r = p.add_run("Note : " + text)
    r.italic = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
    r.font.color.rgb = GRAY

def bold_inline(p, bold_text, normal_text=''):
    """Ajoute un fragment gras + texte normal dans un paragraphe existant."""
    rb = p.add_run(bold_text)
    rb.bold = True; rb.font.name = 'Calibri'; rb.font.size = Pt(11)
    if normal_text:
        rn = p.add_run(normal_text)
        rn.font.name = 'Calibri'; rn.font.size = Pt(11)

def mixed_line(bold_part, normal_part, before=0, after=4, indent=False):
    """Ligne avec une partie en gras + partie normale."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    bold_inline(p, bold_part, normal_part)

def spacer(pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)


# ─────────────────────────────────────────────────────────────────────────────
#  PAGE DE TITRE
# ─────────────────────────────────────────────────────────────────────────────

title_line("StuGo CO2 Explorer")
subtitle("Guide d'utilisation  ·  Version 7.0")
hrule()

body(
    "Ce guide explique comment utiliser le logiciel StuGo CO2 Explorer. "
    "Il couvre chaque écran du programme, de l'import de fichiers jusqu'à l'export "
    "des graphiques. Vous n'avez besoin d'aucune connaissance technique particulière.",
    before=12, after=8
)

body(
    "Le logiciel sert à analyser les émissions de CO2 liées aux déplacements "
    "des étudiants en mobilité internationale. Il lit des fichiers Excel au format "
    "StuGo ainsi que des exports CSV produits par ce même logiciel, "
    "les affiche sous forme de tableau, et génère des graphiques.",
    after=6
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 0 : PRÉSENTATION GÉNÉRALE
# ─────────────────────────────────────────────────────────────────────────────

chapter("0", "Présentation générale")

body(
    "Quand vous ouvrez le logiciel, vous voyez une barre de navigation sur le côté gauche "
    "avec plusieurs sections. Voici à quoi sert chacune :",
    after=6
)

bullet("Accueil  —  tableau de bord avec les chiffres clés de la session")
bullet("Import  —  charger vos fichiers Excel")
bullet("Tableau  —  visualiser et trier les données sous forme de liste")
bullet("Graphique  —  créer des graphiques à partir des données")
bullet("Comparaison  —  comparer plusieurs fichiers côte à côte")
bullet("Paramètres  —  changer l'apparence et les préférences")
bullet("Aide  —  rappel des formats de fichiers attendus")

spacer()
body(
    "Trois de ces sections (Tableau, Graphique, Comparaison) disposent aussi d'un "
    "panneau de filtres qui s'affiche à droite de la barre de navigation. "
    "Ce panneau permet de restreindre les données affichées sans modifier les fichiers.",
    after=6
)

body(
    "Les zones CO2 utilisées dans le logiciel correspondent à des seuils d'émission "
    "par voyage :",
    after=4
)
bullet("Zone 1  —  moins de 0,5 tCO2e  (voyages courts ou peu polluants)")
bullet("Zone 2  —  de 0,5 à 1 tCO2e")
bullet("Zone 3  —  de 1 à 2 tCO2e")
bullet("Zone 4  —  de 2 à 3 tCO2e")
bullet("Zone 5  —  plus de 3 tCO2e  (voyages longs, fortement émetteurs)")


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 1 : ACCUEIL
# ─────────────────────────────────────────────────────────────────────────────

chapter("1", "L'écran d'accueil")

body(
    "L'accueil s'affiche au démarrage. Il résume l'état de la session en cours "
    "et donne un accès rapide aux autres sections.",
    after=6
)

section("Les statistiques de session")
body(
    "Six chiffres sont affichés en haut de l'écran. Ils se mettent à jour "
    "automatiquement dès que vous importez ou retirez des fichiers :",
    after=4
)
bullet("Étudiants  —  nombre total d'étudiants dans les données chargées")
bullet("Total tCO2e  —  émissions cumulées de tous les voyages")
bullet("Destinations  —  nombre de pays de destination distincts")
bullet("Feuilles importées  —  nombre de feuilles Excel chargées")
bullet("Lignes de données  —  nombre total de lignes dans le tableau")
bullet("Exports effectués  —  nombre de fichiers exportés depuis le démarrage")

section("L'horloge")
body(
    "En haut à droite, une petite horloge affiche l'heure et la date du jour en temps réel. "
    "Elle n'a pas d'autre fonction.",
    after=6
)

section("Les boutons de navigation rapide")
body(
    "Quatre boutons sous le titre permettent d'aller directement à Import, Tableau, "
    "Graphique ou Comparaison sans passer par la barre latérale.",
    after=6
)

section("Le guide d'utilisation intégré")
body(
    "En bas à droite de l'accueil, un encadré rappelle les cinq étapes principales "
    "d'utilisation du logiciel. C'est un rappel visuel, pas interactif.",
    after=6
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 2 : IMPORT
# ─────────────────────────────────────────────────────────────────────────────

chapter("2", "Importer des fichiers")

body(
    "Avant de pouvoir utiliser le Tableau, le Graphique ou la Comparaison, "
    "vous devez d'abord charger au moins un fichier Excel. C'est le point de départ.",
    after=6
)

section("Comment ajouter des fichiers")
body(
    "Deux méthodes sont disponibles, au choix :",
    after=4
)
bullet(
    "Cliquer sur le bouton Ajouter des fichiers (en haut à droite de l'écran Import) "
    "— une fenêtre s'ouvre pour sélectionner un ou plusieurs fichiers .xlsx, .xls ou .csv"
)
bullet(
    "Glisser-déposer un ou plusieurs fichiers directement dans la zone pointillée "
    "qui occupe le haut de l'écran"
)
note(
    "Vous pouvez sélectionner plusieurs fichiers à la fois dans la fenêtre de sélection "
    "en maintenant Ctrl ou Shift tout en cliquant."
)

section("Import de fichiers CSV (nouveau en v7.0)")
body(
    "Depuis la version 7.0, le logiciel accepte également les fichiers .csv "
    "exportés par StuGo CO2 Explorer lui-même (bouton Exporter CSV dans l'écran Tableau). "
    "Attention : seuls les exports produits par ce logiciel sont acceptés. "
    "Un fichier CSV quelconque, issu d'Excel ou d'une autre source, ne sera pas chargé.",
    after=4
)
body("Le fichier CSV doit contenir les colonnes suivantes :", after=4)
bullet("sheet_id  —  identifiant de la feuille source")
bullet("zone  —  numéro de zone CO2")
bullet("zone_label  —  libellé de la zone")
bullet("pays  —  pays de destination")
bullet("tco2e_par_voyage  —  émissions par trajet")
bullet("nb_etudiants  —  nombre d'étudiants")
bullet("total_tco2  —  émissions totales")
bullet("fichier  —  nom du fichier source d'origine")
note(
    "Si une colonne est manquante, le logiciel affiche un message d'erreur clair "
    "indiquant quelle colonne est absente. Le fichier n'est alors pas chargé."
)

section("La liste des fichiers importés")
body(
    "Chaque fichier ajouté apparaît sous forme de carte avec :",
    after=4
)
bullet("Le nom de la feuille Excel")
bullet("Le nom du fichier source")
bullet("Un résumé rapide : nombre d'étudiants, total tCO2e, nombre de pays")

body(
    "Chaque carte a une couleur distincte sur le bord gauche pour faciliter "
    "l'identification dans les autres écrans.",
    after=6
)

section("Retirer un fichier")
body(
    "Pour retirer un fichier de la liste, cliquez sur le petit bouton × à droite "
    "de sa carte. Une confirmation vous est demandée avant la suppression.",
    after=4
)
body(
    "Pour tout retirer d'un coup, utilisez le bouton rouge Tout supprimer "
    "en haut de la liste. Une confirmation est également demandée.",
    after=6
)

note(
    "Retirer un fichier ne supprime pas le fichier de votre ordinateur. "
    "Il est simplement retiré de la session en cours."
)

section("Détection automatique des doublons CSV / Excel")
body(
    "Lorsque vous importez un fichier CSV exporté par le logiciel en même temps qu'un fichier Excel "
    "dont il est issu, le logiciel détecte automatiquement qu'il s'agit des mêmes données. "
    "Un message s'affiche pour vous avertir :",
    after=4
)
bullet(
    "Non (recommandé) — ignore le CSV et l'Excel en double. "
    "Les fichiers non-concernés par le doublon sont importés normalement."
)
bullet(
    "Oui — importe les deux quand même (les données apparaîtront deux fois dans le tableau)."
)
note(
    "Cette détection fonctionne grâce à la colonne 'fichier' présente dans les CSV exportés, "
    "qui conserve le nom du fichier Excel d'origine. Le logiciel compare cette information "
    "avec les fichiers Excel déjà chargés."
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 3 : LES FILTRES
# ─────────────────────────────────────────────────────────────────────────────

chapter("3", "Utiliser les filtres")

body(
    "Le panneau de filtres est disponible dans les sections Tableau, Graphique "
    "et Comparaison. Il se trouve à droite de la barre de navigation. "
    "Les filtres s'appliquent instantanément dès que vous les modifiez.",
    after=6
)

section("Filtrer par faculté")
body(
    "Une liste de cases à cocher apparaît avec toutes les facultés présentes dans "
    "vos données. Décochez celles que vous ne souhaitez pas voir. "
    "Par défaut, toutes sont cochées.",
    after=6
)

section("Filtrer par zone CO2")
body(
    "Cinq cases correspondent aux cinq zones (Zone 1 à Zone 5). "
    "Décochez une zone pour masquer toutes les données de cette tranche d'émission.",
    after=6
)

section("Filtrer par pays")
body(
    "Un champ texte permet de rechercher un pays par son nom. "
    "Tapez les premières lettres et seules les lignes dont le pays contient "
    "ce texte seront affichées.",
    after=6
)

section("Filtrer par nombre d'étudiants")
body(
    "Deux champs Min et Max permettent de définir une plage. "
    "Seules les lignes avec un nombre d'étudiants compris entre ces deux valeurs "
    "seront affichées.",
    after=4
)
body(
    "La case Masquer les 0 étudiant permet de cacher les lignes où le nombre "
    "d'étudiants est nul.",
    after=6
)

section("Filtrer par total tCO2e")
body(
    "Même principe que pour les étudiants : définissez un Min et un Max "
    "pour ne garder que les lignes dont le total tCO2e se situe dans cette plage.",
    after=6
)

section("Réinitialiser les filtres")
body(
    "Le bouton Réinit. en haut du panneau de filtres remet tous les filtres "
    "à leur état par défaut : tout coché, aucune plage.",
    after=6
)

note(
    "Les filtres ne modifient pas vos fichiers. Ils changent seulement ce qui est "
    "affiché à l'écran. Si vous passez d'un écran à l'autre, les filtres restent actifs."
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 4 : TABLEAU
# ─────────────────────────────────────────────────────────────────────────────

chapter("4", "Le Tableau")

body(
    "L'écran Tableau affiche toutes les données importées sous forme de liste. "
    "C'est l'endroit où consulter les données en détail, les trier et les exporter.",
    after=6
)

section("La barre de statistiques")
body(
    "En haut du tableau, une barre affiche en permanence les chiffres clés "
    "calculés sur les données actuellement filtrées :",
    after=4
)
bullet("Lignes  —  nombre de lignes affichées")
bullet("Étudiants  —  somme du nombre d'étudiants")
bullet("TCO2e total  —  somme des émissions")
bullet("Pays  —  nombre de pays distincts")
bullet("Moy. tCO2e/voy.  —  moyenne des émissions par voyage")

section("Les colonnes du tableau")
body("Le tableau contient les colonnes suivantes :", after=4)
bullet("Faculté  —  nom de la feuille source")
bullet("Zone N°  —  numéro de la zone CO2 (1 à 5)")
bullet("Zone  —  libellé de la zone (avec les seuils)")
bullet("Pays  —  pays de destination")
bullet("tCO2e / voyage  —  émissions par trajet")
bullet("Étudiants  —  nombre d'étudiants concernés")
bullet("Total TCO2  —  émissions totales pour cette ligne")
bullet("Fichier  —  nom du fichier source")

section("Trier les données")
body(
    "Deux façons de trier le tableau :",
    after=4
)
bullet(
    "Cliquer directement sur l'en-tête d'une colonne pour trier par cette colonne. "
    "Cliquer une deuxième fois inverse l'ordre."
)
bullet(
    "Utiliser la liste déroulante Tri en haut à droite pour choisir une colonne, "
    "et le bouton Asc / Desc pour changer le sens."
)

section("Exporter en CSV")
body(
    "Le bouton Exporter CSV en haut à droite génère un fichier .csv "
    "avec les données telles qu'elles sont affichées à ce moment (filtres appliqués). "
    "Une fenêtre vous demande où enregistrer le fichier.",
    after=4
)
note(
    "Le bouton Exporter CSV est grisé (inactif) tant qu'aucun fichier n'est importé. "
    "Il s'active automatiquement dès que des données sont chargées."
)
note(
    "Le fichier CSV peut être ouvert avec Excel ou tout autre tableur. "
    "Seules les données visibles sont exportées — les lignes masquées par les filtres "
    "ne sont pas incluses."
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 5 : GRAPHIQUE
# ─────────────────────────────────────────────────────────────────────────────

chapter("5", "L'écran Graphique")

body(
    "L'écran Graphique génère un graphique à partir des données filtrées. "
    "Il se met à jour en temps réel quand vous changez les options.",
    after=6
)

section("Choisir le type de graphique")
body("Six types sont disponibles dans la liste Type :", after=4)
bullet("Barres verticales  —  un graphique en colonnes classique")
bullet("Barres horizontales  —  les mêmes colonnes, mais à l'horizontale")
bullet("Camembert  —  graphique circulaire avec parts proportionnelles")
bullet("Donut  —  comme le camembert, avec un trou au centre")
bullet("Aires empilées  —  courbes avec surfaces colorées empilées")
bullet("Treemap  —  rectangles proportionnels aux valeurs (vue d'ensemble)")

section("Mode 2D et mode 3D")
body(
    "Les boutons 2D et 3D en haut à gauche basculent entre les deux modes. "
    "Le mode 3D ajoute un effet de profondeur et des contrôles de rotation.",
    after=4
)
body(
    "En mode 3D, une barre supplémentaire apparaît avec :",
    after=4
)
bullet("Élévation  —  curseur pour incliner la vue de haut en bas (de -90° à 90°)")
bullet("Azimut  —  curseur pour faire pivoter la vue de gauche à droite (-180° à 180°)")
bullet("Reset  —  bouton pour revenir à l'angle par défaut")

note(
    "Tous les types de graphiques ne sont pas disponibles en 3D. "
    "Le Treemap, par exemple, reste en 2D même si le mode 3D est activé. "
    "Le logiciel l'indique par un message."
)

section("Configurer les axes")
body("Trois listes permettent de personnaliser ce que le graphique affiche :", after=4)
mixed_line("Axe X : ", "ce qui est représenté en catégories — Faculté, Zone ou Pays", indent=True)
mixed_line("Valeur : ", "la grandeur mesurée — Total tCO2e, Nombre d'étudiants ou tCO2e par voyage", indent=True)
mixed_line("Couleur : ", "comment les données sont colorées — par Faculté, par Zone, ou sans distinction", indent=True)
spacer(4)
mixed_line("Top pays : ", "si l'axe X est Pays, ce chiffre limite le nombre de pays affichés (entre 3 et 50)", indent=True)

section("La barre de navigation du graphique")
body(
    "Sous les contrôles, une petite barre d'outils matplotlib apparaît au-dessus du graphique. "
    "Elle permet :",
    after=4
)
bullet("De zoomer en traçant un rectangle sur le graphique")
bullet("De se déplacer dans le graphique en maintenant le clic")
bullet("De revenir à la vue d'origine")
bullet("De sauvegarder directement le graphique en image")

section("Exporter le graphique")
body(
    "Le bouton Exporter PNG en haut de l'écran ouvre une fenêtre d'enregistrement. "
    "Trois formats sont disponibles :",
    after=4
)
bullet("PNG  —  image bitmap, idéale pour les présentations et documents Word")
bullet("PDF  —  document vectoriel, parfait pour l'impression")
bullet("SVG  —  graphique vectoriel, modifiable dans des outils comme Inkscape")


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 6 : COMPARAISON
# ─────────────────────────────────────────────────────────────────────────────

chapter("6", "L'écran Comparaison")

body(
    "L'écran Comparaison permet de visualiser plusieurs fichiers en même temps "
    "et de les comparer graphiquement. C'est utile quand vous avez importé "
    "des données de plusieurs années ou de plusieurs établissements.",
    after=6
)

section("Le panneau de sélection des fichiers (côté gauche)")
body(
    "Sur la gauche de l'écran, une liste affiche tous les fichiers importés. "
    "Chaque fichier a une case à cocher. "
    "Seuls les fichiers cochés apparaissent dans le graphique.",
    after=4
)
body("Trois outils facilitent la sélection :", after=4)
bullet("Le bouton ✓  —  coche tous les fichiers d'un coup")
bullet("Le bouton ✕  —  décoche tous les fichiers d'un coup")
bullet(
    "La barre de recherche (en haut de la liste) "
    "—  tapez quelques lettres pour filtrer les noms de fichiers affichés"
)
body(
    "Le menu Tri permet de réordonner la liste des fichiers selon :",
    after=4
)
bullet("Nom  —  ordre alphabétique")
bullet("CO2 ↓ et CO2 ↑  —  du plus au moins émetteur (ou inversement)")
bullet("Étud. ↓ et Étud. ↑  —  du plus grand au plus petit nombre d'étudiants")
spacer(4)
body(
    "En bas du panneau, un compteur indique combien de fichiers sont sélectionnés "
    "sur le total.",
    after=6
)

section("Les modes d'affichage")
body(
    "La liste déroulante Mode, dans la barre de contrôles, propose quatre façons "
    "de présenter les fichiers sélectionnés :",
    after=4
)
mixed_line(
    "Grille par fichier  —  ",
    "un graphique séparé pour chaque fichier, disposés en grille. "
    "Le nombre de colonnes de la grille se règle avec le compteur Cols (1 à 5)."
)
mixed_line(
    "Graphe unifié  —  ",
    "toutes les données de tous les fichiers sélectionnés sont fusionnées "
    "dans un seul graphique. Utile pour une vue d'ensemble."
)
mixed_line(
    "Côte à côte  —  ",
    "barres groupées : pour chaque valeur de l'axe X, "
    "une barre par fichier côte à côte. Permet de comparer directement."
    " (disponible uniquement avec les types Barres verticales ou Barres horizontales)"
)
mixed_line(
    "Empilé par fichier  —  ",
    "barres empilées : chaque fichier est représenté par une barre, "
    "dont les segments correspondent aux zones ou facultés. "
    "Montre la composition de chaque fichier."
    " (disponible uniquement avec les types Barres verticales ou Barres horizontales)"
)

section("Les contrôles du graphique")
body(
    "Les mêmes contrôles que dans l'écran Graphique sont disponibles ici : "
    "Type de graphique, 2D/3D, Axe X, Valeur, Couleur, Top pays. "
    "Référez-vous au chapitre 5 pour leur description.",
    after=6
)

section("Exporter les graphiques")
body(
    "Le bouton PNG en haut à droite exporte tous les graphiques actuellement affichés "
    "dans un dossier de votre choix. "
    "Un fichier PNG est créé pour chaque graphique, nommé d'après le fichier source.",
    after=6
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 7 : PARAMÈTRES
# ─────────────────────────────────────────────────────────────────────────────

chapter("7", "Les Paramètres")

body(
    "L'écran Paramètres regroupe les options d'apparence et de comportement du logiciel. "
    "Les modifications prennent effet immédiatement, sans redémarrage.",
    after=6
)

section("Changer le thème de couleurs")
body(
    "La section Thème affiche sept thèmes prédéfinis, chacun représenté "
    "par un bouton aux couleurs du thème :",
    after=4
)
bullet("Arctic Ink (défaut)  —  fond très sombre, accent turquoise")
bullet("Sombre  —  fond sombre type GitHub, accent vert")
bullet("Clair Pro  —  fond blanc, accents verts, adapté à un usage en plein jour")
bullet("Violet Nuit  —  fond noir profond, accents violets")
bullet("Bleu Marine  —  fond bleu nuit, accents bleu clair")
bullet("Solarized  —  fond teal sombre, palette Solarized classique")
bullet("Sable Clair  —  fond beige, accents or et terre")
spacer(4)
body(
    "Cliquez sur un thème pour l'appliquer instantanément. "
    "Le thème sélectionné est mémorisé et sera rechargé au prochain démarrage.",
    after=6
)

section("Créer un preset de couleurs personnalisé")
body(
    "La section Preset personnalisé permet de créer votre propre thème de couleurs. "
    "Choisissez d'abord le mode dans la liste Mode :",
    after=4
)
mixed_line(
    "Minimal (2 clés)  —  ",
    "vous choisissez seulement la couleur de fond et la couleur d'accent. "
    "Toutes les autres couleurs sont calculées automatiquement à partir de ces deux-là."
)
mixed_line(
    "Essentiel (6 clés)  —  ",
    "vous définissez 6 couleurs principales. "
    "Les 12 couleurs restantes sont déduites automatiquement."
)
mixed_line(
    "Complet (18 clés)  —  ",
    "vous contrôlez les 18 couleurs de la palette une par une."
)
spacer(6)
body("Pour changer une couleur :", after=4)
bullet("Cliquez sur le bouton coloré correspondant")
bullet("La fenêtre de sélection de couleur s'ouvre")
bullet("Choisissez la couleur souhaitée et validez")
bullet("Le bouton se met à jour avec la nouvelle couleur")
spacer(4)
body("Trois boutons d'action sont disponibles :", after=4)
bullet(
    "Aperçu  —  applique le preset temporairement pour voir le résultat. "
    "Si vous n'aimez pas, changez de thème dans la section Thème pour annuler."
)
bullet(
    "Sauvegarder (.json)  —  enregistre le preset dans un fichier JSON "
    "dans votre dossier personnel. Vous pouvez nommer le preset en haut de la section."
)
bullet(
    "Charger  —  importe un ou plusieurs fichiers JSON de preset "
    "précédemment sauvegardés. Ils apparaissent ensuite dans la section Thème."
)

section("Changer le curseur de la souris")
body(
    "La section Curseur propose huit styles de curseur prédéfinis pour l'intérieur "
    "du logiciel :",
    after=4
)
bullet("Défaut système  —  le curseur habituel de Windows")
bullet("Flèche standard, Main pointeur, Curseur texte, Croix de visée, etc.")
bullet(
    "Image personnalisée  —  utilisez votre propre image (.png, .jpg, .jpeg, .ico)."
)
spacer(4)
body(
    "Pour utiliser une image personnalisée, cliquez sur le bouton Parcourir "
    "(toujours accessible). Le logiciel bascule automatiquement sur Image personnalisée "
    "dans la liste déroulante, puis ouvre une fenêtre de sélection de fichier. "
    "Choisissez votre image — elle est copiée dans votre dossier personnel "
    "et appliquée immédiatement.",
    after=4
)
note("Le curseur choisi est mémorisé et rechargé au prochain démarrage.")
note(
    "Si l'image sélectionnée est invalide ou illisible, un message d'erreur s'affiche "
    "et le curseur précédent est conservé. "
    "Pour revenir au curseur système, cliquez sur Restaurer le curseur par défaut "
    "ou sélectionnez Défaut système dans la liste."
)

section("Sauvegarde de session et fermeture du logiciel")
body(
    "Le logiciel surveille vos activités (import de fichiers, suppressions) "
    "et vous avertit si vous essayez de fermer la fenêtre sans avoir sauvegardé. "
    "Un message apparaît avec trois choix :",
    after=4
)
bullet(
    "Sauvegarder et quitter (recommandé)  —  enregistre la session puis ferme le logiciel. "
    "La prochaine fois que vous ouvrez le logiciel, vous pourrez reprendre exactement "
    "là où vous en étiez."
)
bullet(
    "Quitter sans sauvegarder  —  ferme le logiciel sans rien enregistrer. "
    "Les fichiers chargés et les actions du journal seront perdus."
)
bullet(
    "Annuler  —  revient au logiciel sans le fermer."
)
note(
    "Si la session a déjà été sauvegardée et qu'aucun changement n'a été fait depuis, "
    "le logiciel se ferme directement sans afficher ce message."
)
spacer(4)
body(
    "Lorsque vous cliquez sur Sauvegarder la session dans le Journal d'activité, "
    "et qu'une session précédente existe déjà, le logiciel vous demande "
    "si vous souhaitez l'écraser. "
    "Répondez Oui pour remplacer l'ancienne sauvegarde, Non pour annuler.",
    after=6
)

section("Effacer les données de session")
body(
    "Le bouton Effacer les données de session supprime les informations "
    "de la session en cours (derniers fichiers chargés, état des filtres, etc.). "
    "Une confirmation est demandée. "
    "Cela ne supprime pas vos presets de couleurs ni vos préférences de thème.",
    after=6
)


# ─────────────────────────────────────────────────────────────────────────────
#  CHAPITRE 8 : QUESTIONS FRÉQUENTES
# ─────────────────────────────────────────────────────────────────────────────

chapter("8", "Questions fréquentes")

questions = [
    (
        "Le graphique est vide ou ne s'affiche pas",
        "Vérifiez qu'au moins un fichier est importé et qu'aucun filtre ne masque "
        "toutes les données. Dans l'écran Comparaison, vérifiez qu'au moins un fichier "
        "est coché dans la liste de gauche."
    ),
    (
        "Les statistiques sur l'accueil affichent 0",
        "C'est normal si aucun fichier n'est importé. "
        "Allez dans Import et chargez vos fichiers Excel."
    ),
    (
        "Je ne vois pas toutes mes données dans le tableau",
        "Un ou plusieurs filtres sont probablement actifs. "
        "Cliquez sur Réinit. dans le panneau de filtres pour tout remettre à zéro."
    ),
    (
        "Mon fichier Excel ne se charge pas",
        "Vérifiez que le fichier est bien au format .xlsx ou .xls "
        "et qu'il correspond au format StuGo CO2 Explorer. "
        "Consultez l'onglet Aide dans le logiciel pour les détails du format attendu."
    ),
    (
        "Comment revenir au thème d'origine après un aperçu de preset ?",
        "Allez dans Paramètres, section Thème, et cliquez sur Arctic Ink (défaut) "
        "ou sur le thème que vous utilisiez."
    ),
    (
        "Le mode 3D affiche un message sur l'incompatibilité",
        "Certains types de graphiques n'ont pas de rendu 3D (Treemap notamment). "
        "Le logiciel reste en 2D dans ce cas. Essayez un autre type comme Barres verticales."
    ),
    (
        "Mes presets personnalisés ont disparu après une réinstallation",
        "Les presets sont dans votre dossier personnel. "
        "Appuyez sur Windows + R, tapez %LOCALAPPDATA%\\StuGoCO2 et appuyez sur Entrée. "
        "Vos fichiers .json se trouvent dans le sous-dossier PresetColors."
    ),
    (
        "Comment exporter toutes les données en même temps ?",
        "Dans l'écran Tableau, le bouton Exporter CSV exporte toutes les données visibles. "
        "Dans l'écran Comparaison, le bouton PNG exporte tous les graphiques affichés "
        "en une seule opération dans un dossier."
    ),
    (
        "Un message me demande si je veux sauvegarder avant de quitter",
        "C'est normal. Le logiciel détecte que vous avez importé des fichiers ou effectué "
        "des actions depuis le dernier démarrage. Cliquez sur Sauvegarder et quitter pour "
        "conserver votre travail, ou Quitter sans sauvegarder si vous n'en avez pas besoin."
    ),
    (
        "Un avertissement apparaît quand je sauvegarde la session",
        "Une session précédente existe déjà. Le logiciel vous demande si vous souhaitez "
        "l'écraser. Répondez Oui pour remplacer l'ancienne, Non pour annuler la nouvelle "
        "sauvegarde et conserver l'ancienne."
    ),
    (
        "Mon CSV ne se charge pas ou est refusé à l'import",
        "Seuls les CSV exportés par StuGo CO2 Explorer via le bouton Exporter CSV sont acceptés. "
        "Un CSV créé manuellement ou depuis Excel sera refusé. "
        "Le message d'erreur indique quelles colonnes sont manquantes."
    ),
    (
        "Un doublon a été détecté entre un Excel et un CSV",
        "Le logiciel a reconnu que le CSV est un export du fichier Excel déjà chargé. "
        "Cliquez sur Non pour ignorer le doublon (recommandé) ou Oui pour importer les deux."
    ),
    (
        "Où trouver les journaux d'erreurs du logiciel ?",
        "Les erreurs techniques sont enregistrées dans votre dossier personnel. "
        "Appuyez sur Windows + R, tapez %LOCALAPPDATA%\\StuGoCO2\\logs et appuyez sur Entrée. "
        "Les fichiers sont nommés log-AAAA-MM-JJ.txt (un par jour). "
        "Vous pouvez les ouvrir avec le Bloc-notes."
    ),
    (
        "Le bouton Parcourir (image curseur) ne fait rien quand je clique dessus",
        "Dans la version installée, le logiciel utilise son propre sélecteur de fichier "
        "pour garantir le bon fonctionnement. Si la fenêtre ne s'ouvre pas, "
        "consultez le fichier log du jour dans %LOCALAPPDATA%\\StuGoCO2\\logs\\ "
        "pour identifier le message d'erreur précis."
    ),
    (
        "J'ai importé une image curseur mais le curseur n'a pas changé",
        "Vérifiez que l'image est bien au format .png, .jpg ou .ico. "
        "Le logiciel valide l'image avant de l'accepter — un message d'erreur s'affiche "
        "si le fichier n'est pas lisible. "
        "Vous pouvez aussi cliquer sur Restaurer le curseur par défaut, "
        "puis réessayer l'import."
    ),
]

for q, a in questions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(q)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11)
    body(a, before=0, after=6, indent=True)


# ─────────────────────────────────────────────────────────────────────────────
#  PIED DE PAGE
# ─────────────────────────────────────────────────────────────────────────────

spacer(12)
hrule()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_before = Pt(6)
r = pf.add_run("StuGo CO2 Explorer v7.0  —  Guide d'utilisation")
r.font.name = 'Calibri'; r.font.size = Pt(9)
r.font.color.rgb = GRAY


doc.save('GUIDE_UTILISATEUR.docx')
print('Document cree : GUIDE_UTILISATEUR.docx')
