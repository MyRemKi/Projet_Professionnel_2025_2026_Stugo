"""Generate the technical report for StuGo CO2 Explorer."""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date

OUTPUT = "Rapport_Technique_StuGoCO2.docx"
BLUE = RGBColor(0x1A, 0x56, 0x9A)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
doc.core_properties.author = "Remi Kalkan"
doc.core_properties.last_modified_by = "Remi Kalkan"
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = BLUE


def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.style.font.color.rgb = BLUE
    return p


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = BLUE
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


title("StuGo CO2 Explorer")
subtitle("Rapport technique : architecture, méthodes, conformité et remise du code source")
subtitle(f"Version analysée : v6.1 responsive / documentation de build v7.0  |  {date.today().strftime('%d/%m/%Y')}")

heading("1. Synthèse exécutive")
para("StuGo CO2 Explorer est une application de bureau Python qui fonctionne hors ligne. Elle importe des fichiers Excel et CSV, normalise les données de mobilité étudiante, applique des filtres, affiche des tableaux et rend des graphiques 2D/3D. Aucun appel réseau applicatif n'a été identifié dans le périmètre du code source.")
para("Les correctifs demandés concernant l'emplacement des données utilisateur sont présents dans le code et dans l'installeur : les préférences, sessions et logs de l'exécutable sont écrits dans le profil de l'utilisateur sous %LOCALAPPDATA%\\StuGoCO2, et les clés d'installation sont créées sous HKEY_CURRENT_USER (HKCU).")
para("Conclusion : l'architecture met effectivement en oeuvre plusieurs design patterns classiques. La réponse aux exigences du service informatique est favorable pour l'exécutable installé ; la seule nuance est que le mode développeur écrit volontairement ses logs dans le dossier du projet pour faciliter le diagnostic local.")

heading("2. Répartition du code")
table(["Couche", "Répertoires / modules", "Responsabilité"], [
    ("Entrée", "main.py, app/bootstrap.py", "Initialiser QApplication, DPI, thème et fenêtre principale."),
    ("Application", "app/commands, app/events, app/services", "Orchestrer les cas d'usage, commandes, événements et services."),
    ("Domaine", "domain/repositories, domain/value_objects", "Définir les contrats et objets métier indépendants de Qt et du stockage."),
    ("Infrastructure", "infrastructure/extractors, models, persistence", "Lire Excel/CSV, adapter pandas à Qt et persister les données JSON."),
    ("Présentation", "presentation/pages, sidebar, widgets, theme", "Construire l'interface PyQt6 et relier les interactions aux services."),
    ("Rendu", "rendering/factory, strategies, axes, geometry", "Sélectionner et exécuter les rendus graphiques 2D/3D."),
    ("Partagé", "shared/constants, paths, scaling, logging", "Centraliser chemins, constantes, mise à l'échelle et journaux."),
])

heading("3. Classes et méthodes principales")
para("La liste ci-dessous couvre les classes qui portent les responsabilités architecturales. Les fonctions utilitaires privées des extracteurs et du rendu sont regroupées par responsabilité afin de rester lisibles.")
table(["Classe / module", "Méthodes principales", "Rôle"], [
    ("ChartConfig", "default, with_type, with_3d, as_2d", "Objet-valeur immuable décrivant une configuration de graphique."),
    ("ColorValue", "from_rgb, darkened, lightened, hue_shifted, to_rgb, contrast_text", "Encapsuler et transformer une couleur."),
    ("SessionRepository", "instance, load_session, save_session, clear_session, get_session_dir", "Persister la session en JSON dans AppData."),
    ("PreferencesRepository", "instance, load, save", "Persister les préférences utilisateur en JSON."),
    ("ExtractorFactory", "get, register, supported", "Choisir l'extracteur selon l'extension."),
    ("ExcelExtractor / CsvExtractor", "extract_all_sheets / extract_from_csv", "Valider, lire et normaliser les fichiers d'entrée."),
    ("PandasModel", "rowCount, columnCount, data, headerData, sort, update_df", "Adapter DataFrame à QAbstractTableModel."),
    ("EventBus", "instance, publish, subscribe, unsubscribe", "Diffuser les événements entre composants."),
    ("AppState", "instance, current, transition, is_ready, is_loading", "Gérer les états IDLE, LOADING, READY et ERROR."),
    ("Command / CommandHistory", "execute, undo, can_undo, undo_last, clear", "Encapsuler les actions et gérer l'annulation."),
    ("LoadFilesCommand", "execute, undo, can_undo", "Charger des fichiers et retirer ce chargement."),
    ("ApplyFilterCommand", "execute", "Appliquer les filtres sélectionnés."),
    ("DataService", "instance, load_files, remove_file, filtered_df, reset", "Façade des données, extraction, filtres et modèle."),
    ("ChartService / ExportService", "render / export_png, export_csv", "Façader le rendu et les exports."),
    ("RendererFactory", "get, render", "Sélectionner le renderer 2D ou 3D."),
    ("Renderer2D / Renderer3D", "render_chart, render_chart_3d", "Rendre les graphiques par stratégie."),
    ("Cube3D / Pie3D", "build_*_faces, draw_*", "Construire et dessiner les géométries 3D."),
    ("MainWindow", "_build_ui, _nav, _load_files, _apply_filters, _save_session, closeEvent", "Façade et médiateur de l'interface."),
    ("Pages PyQt6", "update_data, update_df, _refresh, set_app", "Présenter accueil, import, tableau, graphiques, comparaison, paramètres et aide."),
    ("FilterSidebar", "refresh_faculties, apply, reset_filters, get_state, restore_state", "Filtrer et restaurer l'état des filtres."),
    ("ThemeManager / StylesheetBuilder", "apply, register_preset / build", "Gérer les thèmes et produire le QSS."),
    ("FileLogger / ActionLogger", "log, get_logs, clear, to_json, from_json", "Écrire les erreurs sur fichier et conserver l'historique d'actions en mémoire."),
])

heading("4. Diagramme de classes Mermaid")
para("Le diagramme complet est conservé dans architecture.mmd. L'extrait suivant représente les relations structurantes et peut être copié dans Mermaid Live, GitHub ou un éditeur compatible Mermaid.")
mermaid = '''classDiagram\n    class ChartConfig { +with_type() +with_3d() +as_2d() }\n    class ColorValue { +from_rgb() +darkened() +lightened() +to_rgb() }\n    class SessionRepository { +instance() +load_session() +save_session() +clear_session() }\n    class PreferencesRepository { +instance() +load() +save() }\n    class ExtractorFactory { +get(extension) +register() +supported() }\n    class DataService { +instance() +load_files() +remove_file() +filtered_df() +reset() }\n    class EventBus { +publish() +subscribe() +unsubscribe() }\n    class CommandHistory { +execute() +undo_last() +clear() }\n    class RendererFactory { +get(mode_3d) +render() }\n    class MainWindow { +_load_files() +_apply_filters() +_save_session() +closeEvent() }\n    SessionRepository ..|> ISessionRepository\n    PreferencesRepository ..|> IPreferencesRepository\n    DataService --> ExtractorFactory\n    DataService --> PandasModel\n    CommandHistory --> Command\n    LoadFilesCommand --|> Command\n    ApplyFilterCommand --|> Command\n    DataService --> EventBus\n    ChartService --> RendererFactory\n    RendererFactory --> Renderer2D\n    RendererFactory --> Renderer3D\n    MainWindow --> DataService\n    MainWindow --> SessionRepository\n    MainWindow --> EventBus'''
for line in mermaid.splitlines():
    code(line)

heading("5. Méthodes utilisées et flux d'exécution")
para("Import : MainWindow déclenche LoadFilesCommand ; DataService délègue le choix à ExtractorFactory ; l'extracteur valide les feuilles, ignore les modèles et dédoublonne les couples fichier/feuille ; le résultat est normalisé dans pandas puis exposé à PandasModel.")
para("Filtrage : FilterSidebar conserve l'état des facultés et zones. ApplyFilterCommand demande au DataService de produire un DataFrame filtré, qui est propagé aux pages tableau, graphique et comparaison par signaux/événements Qt.")
para("Graphiques : ChartPage transmet la configuration à ChartService. RendererFactory choisit la stratégie 2D ou 3D ; les modules d'axes appliquent le style et Cube3D/Pie3D construisent les géométries spécialisées.")
para("Session et préférences : SessionRepository et PreferencesRepository sérialisent des dictionnaires JSON. MainWindow restaure la session au démarrage et la sauvegarde à la fermeture ; SettingsPage sauvegarde les préférences de thème, curseur et présentation.")

heading("6. Design patterns")
table(["Pattern", "Implémentation constatée", "Évaluation"], [
    ("Singleton", "Repositories, DataService, EventBus, AppState, ThemeManager et loggers via instance().", "Oui. Centralise les services sans multiplier les instances."),
    ("Factory", "ExtractorFactory et RendererFactory.", "Oui. Le code appelant ne dépend pas du choix concret."),
    ("Strategy", "Renderers 2D/3D sélectionnés selon le mode et le type de graphique.", "Oui, surtout dans la couche rendering."),
    ("Command", "Command, LoadFilesCommand, ApplyFilterCommand et CommandHistory.", "Oui. Les actions sont encapsulées et le chargement est réversible."),
    ("Observer / Publish-Subscribe", "EventBus et signaux Qt.", "Oui. Les composants réagissent aux événements sans couplage direct."),
    ("Adapter / MVC", "PandasModel adapte DataFrame à QAbstractTableModel.", "Oui. Adaptation claire entre données et vue Qt."),
    ("Facade", "DataService, ChartService, ExportService et MainWindow.", "Oui. API simplifiée devant plusieurs sous-systèmes."),
    ("Mediator", "MainWindow et ui_mediator.py coordonnent les pages et sidebars.", "Oui. Coordination centralisée de l'interface."),
    ("Template Method", "BasePage fournit la structure commune des pages.", "Présent selon les spécialisations de pages."),
    ("Value Object", "ChartConfig frozen dataclass et ColorValue.", "Oui. Valeurs métier encapsulées, ChartConfig immuable."),
])
para("Réponse : oui, le projet répond aux notions de design patterns, avec des usages identifiables dans le code. Il s'agit d'une combinaison pragmatique de patterns, et non d'une application dogmatique de tous les patterns existants.")

heading("7. Réponse aux demandes du service informatique")
table(["Demande", "Réponse documentée"], [
    ("Ne pas écrire dans Program Files", "L'installeur utilise {localappdata}\\Programs\\StuGo CO2 Explorer. Les données de l'utilisateur utilisent %LOCALAPPDATA%\\StuGoCO2. Aucun chemin applicatif source ne cible Program Files."),
    ("Logs et préférences dans AppData", "En mode frozen/PyInstaller, FileLogger utilise %LOCALAPPDATA%\\StuGoCO2\\logs ; PreferencesRepository utilise %LOCALAPPDATA%\\StuGoCO2\\prefs.json ; SessionRepository utilise session.json au même emplacement."),
    ("Écritures HKEY_CURRENT_USER", "installer.iss contient uniquement Root: HKCU pour InstallPath et Version. Aucune écriture HKLM n'est déclarée."),
    ("Langage exact", "Python 3.10 ou supérieur, interface PyQt6. Le packaging est réalisé avec PyInstaller et l'installation avec Inno Setup 6."),
    ("Démonstration 30 min max", "Créneaux proposés, à confirmer : mardi 01/09/2026 10:00-10:30 ; mercredi 02/09/2026 14:00-14:30 ; jeudi 03/09/2026 10:00-10:30. Démonstration : import, filtres, graphiques, export, session et emplacement des fichiers."),
    ("Code source complet", "Le dépôt remis contient les sources Python, le diagramme architecture.mmd, les scripts de build, l'installeur et la documentation. Les dossiers dist/build sont des artefacts de distribution et non la source de référence."),
    ("Dépendances", "PyQt6 >= 6.4, matplotlib >= 3.7, pandas >= 2.0, openpyxl >= 3.1, squarify >= 0.4. Bibliothèque standard Python : json, os, sys, pathlib/typing, datetime, threading, dataclasses, abc."),
    ("Tests en environnement isolé", "Le fonctionnement hors ligne est compatible avec le périmètre observé : import local, calculs locaux, rendu local et persistance locale. La validation finale doit être réalisée par l'équipe technique sur un poste isolé."),
    ("Antivirus / droits", "L'installation est configurée avec PrivilegesRequired=lowest. Les écritures applicatives sont déplacées dans l'espace utilisateur, ce qui répond au scénario de blocage des dossiers protégés. La signature numérique de l'exécutable reste un sujet distinct."),
])

heading("8. Emplacements et vérification des écritures")
code("%LOCALAPPDATA%\\StuGoCO2\\prefs.json       préférences")
code("%LOCALAPPDATA%\\StuGoCO2\\session.json     session")
code("%LOCALAPPDATA%\\StuGoCO2\\logs\\log-YYYY-MM-DD.txt  logs en exécutable")
code("HKCU\\Software\\StuGo\\StuGo CO2 Explorer\\InstallPath")
code("HKCU\\Software\\StuGo\\StuGo CO2 Explorer\\Version")
para("Point de contrôle : FileLogger choisit le dossier AppData lorsque sys.frozen est vrai, ce qui correspond à l'exécutable PyInstaller installé. Lors d'un lancement avec py main.py, il écrit dans logs/ à la racine du projet ; ce comportement de développement ne doit pas être confondu avec le comportement livré aux utilisateurs.")

heading("9. Limites et recommandations")
bullet("Tester sur un poste standard sans élévation : installation, premier lancement, import, export et fermeture.")
bullet("Vérifier avec Procmon ou un contrôle de profil que le processus ne tente aucune écriture sous Program Files.")
bullet("Vérifier les deux valeurs sous HKCU après installation et confirmer l'absence de clé équivalente sous HKLM.")
bullet("Pour uniformiser le mode développeur avec la production, rendre le chemin AppData également obligatoire lorsque l'application est lancée par Python ; cela supprimerait la seule différence d'emplacement des logs.")
bullet("Envisager la signature Authenticode des exécutables afin de réduire les alertes antivirus liées aux binaires non signés.")

heading("10. Conclusion")
para("Le projet est organisé par responsabilités, dispose d'interfaces de domaine et emploie plusieurs patterns reconnaissables. Les correctifs de stockage utilisateur et de registre sont intégrés dans la version installée : AppData pour les écritures utilisateur et HKCU pour les paramètres d'installation. Le logiciel est conçu pour fonctionner hors ligne et ne nécessite pas de droits administrateur selon la configuration de l'installeur. La validation finale reste à effectuer dans l'environnement isolé de l'équipe technique.")

# Footer
for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("StuGo CO2 Explorer - Rapport technique")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

doc.save(OUTPUT)
print(f"Document cree : {OUTPUT}")
